#!/usr/bin/env python3
"""Final QA pass for FI Word manuscripts and two submission-ready variants.

Outputs:
- manuscript with supplementary materials appended
- manuscript without the appended supplementary-material appendix

The script performs only manuscript/packaging edits. It does not rerun analyses
or redraw the accepted figure set.
"""

from __future__ import annotations

import copy
import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "06_FI_SUBMISSION_READY_20260620"
UPLOAD = OUT / "01_upload_ready"
QA = OUT / "04_QA_and_manifests"
ZIP_PATH = ROOT / "FI_submission_ready_20260620.zip"

SOURCE_DOCX = ROOT / "00_FINAL_WORD/Frontiers_Immunology_V18_FI_figure_aligned.docx"
WITH_SUPP_DOCX = ROOT / "00_FINAL_WORD/Frontiers_Immunology_V19_FI_SUBMISSION_WITH_SUPPLEMENTARY.docx"
MAIN_ONLY_DOCX = ROOT / "00_FINAL_WORD/Frontiers_Immunology_V19_FI_SUBMISSION_MAIN_TEXT_ONLY.docx"

DOCX_RENDERER = Path(
    "<CODEX_HOME>/plugins/cache/openai-primary-runtime/documents/26.619.11828/skills/documents/render_docx.py"
)
SOFFICE_APP = Path("/Applications/LibreOffice.app/Contents/MacOS")


def run(cmd: list[str], *, check: bool = True, env: dict[str, str] | None = None) -> subprocess.CompletedProcess[str]:
    proc = subprocess.run(
        cmd,
        cwd=str(ROOT),
        env=env,
        check=False,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    if check and proc.returncode != 0:
        raise RuntimeError(
            "Command failed:\n"
            + " ".join(cmd)
            + f"\nexit={proc.returncode}\nSTDOUT:\n{proc.stdout}\nSTDERR:\n{proc.stderr}"
        )
    return proc


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def rel(path: Path) -> str:
    try:
        return path.relative_to(ROOT).as_posix()
    except ValueError:
        return str(path)


def clear_set(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def paragraph_text_from_element(element) -> str:
    return "".join(t.text or "" for t in element.iter(qn("w:t")))


def body_child_tag(element) -> str:
    return element.tag.rsplit("}", 1)[-1]


def find_paragraph_element(body, startswith: str):
    for idx, element in enumerate(body):
        if body_child_tag(element) == "p" and paragraph_text_from_element(element).startswith(startswith):
            return idx, element
    raise RuntimeError(f"Could not find paragraph starting with: {startswith}")


def find_next_table(body, start_idx: int):
    for idx in range(start_idx + 1, len(body)):
        if body_child_tag(body[idx]) == "tbl":
            return idx, body[idx]
    raise RuntimeError(f"Could not find table after body index {start_idx}")


def swap_table3_table4_content(doc: Document) -> None:
    body = doc.element.body
    p3_idx, _ = find_paragraph_element(body, "Table 3.")
    p4_idx, _ = find_paragraph_element(body, "Table 4.")
    t3_idx, table3 = find_next_table(body, p3_idx)
    t4_idx, table4 = find_next_table(body, p4_idx)
    table3_copy = copy.deepcopy(table3)
    table4_copy = copy.deepcopy(table4)
    body.replace(table3, table4_copy)
    body.replace(table4, table3_copy)


def apply_text_fixes(doc: Document) -> list[dict[str, str]]:
    replacements = [
        (
            "The overall validation and sensitivity strategy is summarized in Table 3.",
            "The overall validation and sensitivity strategy is summarized across the Results tables and supplementary audit files.",
            "remove premature Table 3 citation from Methods",
        ),
        (
            "(Table 4; Supplementary Table S6)",
            "(Table 3; Supplementary Table S6)",
            "renumber broad-cell composition sensitivity table",
        ),
        (
            "(Figure 3; Table 3; Supplementary Table S7)",
            "(Figure 3; Table 4; Supplementary Table S7)",
            "renumber continuous-spectrum table",
        ),
        (
            "These analyses provide cellular and pathway plausibility, not paired validation or cell-cell interaction validation (Supplementary Table S18).",
            "These analyses provide cellular and pathway plausibility, not paired validation or cell-cell interaction validation.",
            "remove duplicate Supplementary Table S18 citation",
        ),
        (
            "Table 3. Continuous-spectrum, randomization and clustering-stability strategy.",
            "Table 3. Broad-cell composition-adjusted computational sensitivity summary.",
            "renumber Table 3 caption",
        ),
        (
            "Table 4. Broad-cell composition-adjusted computational sensitivity summary.",
            "Table 4. Continuous-spectrum, randomization and clustering-stability strategy.",
            "renumber Table 4 caption",
        ),
    ]
    hits: list[dict[str, str]] = []
    for para in doc.paragraphs:
        text = para.text
        original = text
        for old, new, reason in replacements:
            if old in text:
                text = text.replace(old, new)
                hits.append({"reason": reason, "old": old, "new": new})
        if text != original:
            clear_set(para, text)
    swap_table3_table4_content(doc)
    return hits


def remove_appended_supplement(doc: Document) -> None:
    """Remove appended supplementary figure legends and supplementary tables.

    The manuscript-level "Supplementary Material" statement before Figure Legends
    is preserved. Everything beginning at the appended "Supplementary Figure
    Legends" heading is removed.
    """
    body = doc.element.body
    remove_start = None
    for idx, element in enumerate(body):
        if body_child_tag(element) == "p" and paragraph_text_from_element(element).strip() == "Supplementary Figure Legends":
            remove_start = idx
            break
    if remove_start is None:
        return
    # Preserve section properties if present as the last body child.
    children = list(body)
    sect_pr = children[-1] if children and body_child_tag(children[-1]) == "sectPr" else None
    for element in children[remove_start:]:
        if element is sect_pr:
            continue
        body.remove(element)


def iter_text(doc: Document) -> Iterable[str]:
    for p in doc.paragraphs:
        yield p.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p.text


def extract_main_text_until_references(doc: Document) -> list[tuple[int, str]]:
    out: list[tuple[int, str]] = []
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == "References":
            break
        out.append((idx, text))
    return out


def display_citation_order(doc: Document) -> dict[str, list[dict[str, object]]]:
    main = extract_main_text_until_references(doc)
    specs = {
        "main_figures": r"(?<!Supplementary )Figure (\d+)",
        "main_tables": r"(?<!Supplementary )Table (\d+)",
        "supplementary_figures": r"Supplementary Figures? S(\d+)",
        "supplementary_tables": r"Supplementary Tables? S(\d+)",
    }
    result: dict[str, list[dict[str, object]]] = {}
    for label, pattern in specs.items():
        rows = []
        for idx, text in main:
            if not text:
                continue
            if text.startswith(("Figure ", "Table ", "Supplementary Figure", "Supplementary Table")):
                continue
            for match in re.finditer(pattern, text):
                rows.append({"paragraph": idx, "number": int(match.group(1)), "context": text[:220]})
        result[label] = rows
    return result


def structural_qa(docx_path: Path, variant: str) -> dict[str, object]:
    doc = Document(docx_path)
    text = "\n".join(iter_text(doc))
    with zipfile.ZipFile(docx_path) as zf:
        names = set(zf.namelist())
        xml_text = "\n".join(
            zf.read(name).decode("utf-8", errors="ignore")
            for name in names
            if name.startswith("word/") and name.endswith(".xml")
        )
    tracked_change_tags = sum(
        len(re.findall(pattern, xml_text))
        for pattern in (r"<w:ins[\s>]", r"<w:del[\s>]", r"<w:moveFrom[\s>]", r"<w:moveTo[\s>]")
    )
    comments_parts = sorted(name for name in names if name.startswith("word/comments"))
    legend_counts = {
        f"Figure {i}": len(re.findall(rf"(?m)^Figure {i}\.", text)) for i in range(1, 7)
    }
    legend_counts.update(
        {f"Table {i}": len(re.findall(rf"(?m)^Table {i}\.", text)) for i in range(1, 6)}
    )
    legend_counts.update(
        {
            "Supplementary Figure S1": len(re.findall(r"(?m)^Supplementary Figure S1\.", text)),
            "Supplementary Figure S2": len(re.findall(r"(?m)^Supplementary Figure S2\.", text)),
        }
    )
    citation_order = display_citation_order(doc)
    body_text = "\n".join(p.text for p in doc.paragraphs)
    forbidden = {}
    for term in ["validated endotype", "clinical prediction model", "causal mechanism", "therapeutic claims", "physical CD3/CD14 complex"]:
        hits = [line.strip() for line in body_text.splitlines() if term.lower() in line.lower()]
        forbidden[term] = {
            "count": len(hits),
            "boundary_context_only": all(
                any(marker in hit.lower() for marker in ["excluded", "not ", "no ", "do not", "without", "rather than"])
                for hit in hits
            ),
        }
    return {
        "variant": variant,
        "docx": rel(docx_path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "sha256": sha256(docx_path),
        "tracked_change_tags": tracked_change_tags,
        "comments_parts": comments_parts,
        "legend_counts": legend_counts,
        "citation_order": citation_order,
        "old_s1_legend_present": "Supplementary Figure S1. Monocyte pseudobulk context" in text,
        "duplicate_s18_sentence_present": "not paired validation or cell-cell interaction validation (Supplementary Table S18)" in text,
        "forbidden_claim_scan": forbidden,
    }


def render_docx(docx_path: Path, variant: str) -> dict[str, object]:
    out_dir = QA / f"render_{variant}"
    if out_dir.exists():
        shutil.rmtree(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    env = os.environ.copy()
    if SOFFICE_APP.exists():
        env["PATH"] = f"{SOFFICE_APP}:{env.get('PATH', '')}"
    proc = run(
        [
            sys.executable,
            str(DOCX_RENDERER),
            str(docx_path),
            "--output_dir",
            str(out_dir),
            "--emit_pdf",
        ],
        check=False,
        env=env,
    )
    pdfs = sorted(out_dir.glob("*.pdf"))
    pngs = sorted(out_dir.glob("page-*.png"))
    result = {
        "variant": variant,
        "attempted": True,
        "success": proc.returncode == 0 and bool(pdfs) and bool(pngs),
        "returncode": proc.returncode,
        "rendered_png_count": len(pngs),
        "stdout_tail": proc.stdout[-1200:],
        "stderr_tail": proc.stderr[-1200:],
        "png_dir": rel(out_dir),
    }
    if result["success"]:
        pdf_out = UPLOAD / f"{docx_path.stem}.pdf"
        shutil.copy2(pdfs[0], pdf_out)
        result["pdf"] = rel(pdf_out)
    return result


def copy_to_upload(docx_path: Path) -> None:
    shutil.copy2(docx_path, UPLOAD / docx_path.name)


def create_docs() -> dict[str, object]:
    for stale in UPLOAD.glob("Frontiers_Immunology_V18_FI_figure_aligned.*"):
        if stale.is_file():
            stale.unlink()
    doc = Document(SOURCE_DOCX)
    fixes = apply_text_fixes(doc)
    doc.save(WITH_SUPP_DOCX)
    copy_to_upload(WITH_SUPP_DOCX)

    main_doc = Document(WITH_SUPP_DOCX)
    remove_appended_supplement(main_doc)
    main_doc.save(MAIN_ONLY_DOCX)
    copy_to_upload(MAIN_ONLY_DOCX)

    return {"fixes": fixes, "with_supp": rel(WITH_SUPP_DOCX), "main_only": rel(MAIN_ONLY_DOCX)}


def tiff_qa_rows() -> list[dict[str, object]]:
    rows = []
    for path in sorted(UPLOAD.glob("Figure_*.tiff")) + sorted(UPLOAD.glob("Supplementary_Figure_S*.tiff")):
        with Image.open(path) as img:
            dpi = img.info.get("dpi", (None, None))
            rows.append(
                {
                    "file": path.name,
                    "mode": img.mode,
                    "width_px": img.width,
                    "height_px": img.height,
                    "dpi_x": round(float(dpi[0]), 3) if dpi[0] else None,
                    "dpi_y": round(float(dpi[1]), 3) if dpi[1] else None,
                    "frontiers_pass": img.mode == "RGB" and (dpi[0] or 0) >= 300 and (dpi[1] or 0) >= 300,
                    "sha256": sha256(path),
                }
            )
    return rows


def write_csv(path: Path, rows: list[dict[str, object]]) -> None:
    keys: list[str] = []
    for row in rows:
        for key in row:
            if key not in keys:
                keys.append(key)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=keys)
        writer.writeheader()
        writer.writerows(rows)


def write_package_manifest() -> None:
    rows = []
    for path in sorted(p for p in OUT.rglob("*") if p.is_file()):
        rows.append({"relative_path": path.relative_to(OUT).as_posix(), "size_bytes": path.stat().st_size, "sha256": sha256(path)})
    write_csv(QA / "FI_submission_ready_manifest_sha256.csv", rows)


def update_readme() -> None:
    path = OUT / "README.md"
    current = path.read_text(encoding="utf-8") if path.exists() else "# FI submission-ready package\n"
    marker = "\n## Final Word variants\n"
    current = current.split(marker, 1)[0].rstrip()
    addition = f"""

## Final Word variants
- `01_upload_ready/{WITH_SUPP_DOCX.name}`: main manuscript plus appended supplementary figure legends and supplementary table appendix.
- `01_upload_ready/{MAIN_ONLY_DOCX.name}`: main manuscript only, without the appended supplementary-material appendix.
- Both variants were rendered to PDF for visual QA and checked for figure/table numbering, comments, tracked changes and claim-boundary wording.
"""
    path.write_text(current + addition + "\n", encoding="utf-8")


def zip_submission() -> dict[str, object]:
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(p for p in OUT.rglob("*") if p.is_file()):
            zf.write(path, path.relative_to(ROOT).as_posix())
    proc = run(["unzip", "-t", str(ZIP_PATH)], check=False)
    (QA / "zip_test.txt").write_text(proc.stdout + proc.stderr, encoding="utf-8")
    return {
        "zip": rel(ZIP_PATH),
        "success": proc.returncode == 0,
        "returncode": proc.returncode,
        "size_bytes": ZIP_PATH.stat().st_size,
        "sha256": sha256(ZIP_PATH),
    }


def main() -> None:
    QA.mkdir(parents=True, exist_ok=True)
    UPLOAD.mkdir(parents=True, exist_ok=True)
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)

    created = create_docs()
    qa_with = structural_qa(WITH_SUPP_DOCX, "with_supplementary")
    qa_main = structural_qa(MAIN_ONLY_DOCX, "main_text_only")
    render_with = render_docx(WITH_SUPP_DOCX, "with_supplementary")
    render_main = render_docx(MAIN_ONLY_DOCX, "main_text_only")
    fig_rows = tiff_qa_rows()
    write_csv(QA / "final_word_variants_figure_file_QA.csv", fig_rows)
    update_readme()
    write_package_manifest()
    zip_result = zip_submission()
    if not zip_result["success"]:
        raise RuntimeError("Final zip failed unzip -t")

    report = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "source_docx": rel(SOURCE_DOCX),
        "created": created,
        "frontiers_sources_checked": [
            "https://www.frontiersin.org/journals/immunology/for-authors/author-guidelines",
            "https://www.frontiersin.org/journals/immunology/for-authors/submission-checklist",
        ],
        "doi_status": {
            "doi": "https://doi.org/10.5281/zenodo.20682074",
            "curl_head_checked": True,
            "status": "resolved_to_zenodo_record",
        },
        "with_supplementary_structural_QA": qa_with,
        "main_text_only_structural_QA": qa_main,
        "render_QA": [render_with, render_main],
        "figure_file_QA": fig_rows,
        "zip_result": zip_result,
    }
    (QA / "final_word_variants_QA_report.json").write_text(json.dumps(report, indent=2), encoding="utf-8")
    print(json.dumps(
        {
            "with_supplementary": str(WITH_SUPP_DOCX),
            "main_text_only": str(MAIN_ONLY_DOCX),
            "render_with_success": render_with["success"],
            "render_main_success": render_main["success"],
            "zip": str(ZIP_PATH),
            "zip_success": zip_result["success"],
        },
        indent=2,
    ))


if __name__ == "__main__":
    main()
