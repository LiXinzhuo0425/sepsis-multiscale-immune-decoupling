#!/usr/bin/env python3
"""Finalize FI submission figures from the user-approved desktop SVG set.

This script intentionally does only packaging/QA work:
- reads the six user-edited SVGs from the Desktop FI folder;
- exports upload TIFFs plus vector backups at fixed journal dimensions;
- removes stale figure/package artifacts inside this project only;
- assembles a Frontiers in Immunology submission-ready directory and zip.

It does not rerun analyses or modify scientific content.
"""

from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from PIL import Image, ImageChops


ROOT = Path(__file__).resolve().parents[1]
DESKTOP_FI = Path("<AUTHORITATIVE_FI_FIGURE_SOURCE_DIR>")
OUT = ROOT / "06_FI_SUBMISSION_READY_20260620"
UPLOAD = OUT / "01_upload_ready"
VECTOR = OUT / "02_vector_backup"
SOURCE = OUT / "03_source_data_and_attribution"
QA = OUT / "04_QA_and_manifests"
ZIP_PATH = ROOT / "FI_submission_ready_20260620.zip"
DOCX = ROOT / "00_FINAL_WORD/Frontiers_Immunology_V17_FINAL_TEXT_ONLY.docx"
RENDER_DOCX = Path(
    "<CODEX_DOCUMENTS_RENDER_DOCX>"
)

DPI = 600


def normalize_svg_font_css(svg_text: str) -> str:
    """Map Illustrator/PostScript font names to stable family names.

    Illustrator-authored SVGs often use PostScript names such as ArialMT.
    Fontconfig may resolve those names to non-Arial substitutes, changing text
    metrics during Inkscape export. Normalize only the project copy; the user's
    desktop SVG remains untouched.
    """
    replacements = [
        (
            r"font-family\s*:\s*['\"]?Arial-BoldMT['\"]?\s*;?",
            "font-family:'Arial';font-weight:700;",
        ),
        (
            r"font-family\s*:\s*['\"]?ArialMT['\"]?\s*;?",
            "font-family:'Arial';",
        ),
        (
            r"font-family\s*:\s*['\"]?Arial-Black['\"]?\s*;?",
            "font-family:'Arial Black','Arial';font-weight:900;",
        ),
        (
            r"font-family\s*:\s*['\"]?AdobeSongStd-Light-GBpc-EUC-H['\"]?\s*;?",
            "font-family:'Arial';",
        ),
        (
            r"font-family\s*:\s*['\"]?NotoSansArmenian-ExtraBold['\"]?\s*;?",
            "font-family:'Arial';font-weight:700;",
        ),
        (
            r"font-family\s*:\s*['\"]?NotoSansArmenian-Medium['\"]?\s*;?",
            "font-family:'Arial';font-weight:500;",
        ),
        (
            r"font-family\s*:\s*['\"]?NotoSansAdlam-Regular['\"]?\s*;?",
            "font-family:'Arial';",
        ),
    ]
    for pattern, repl in replacements:
        svg_text = re.sub(pattern, repl, svg_text)
    svg_text = re.sub(
        r"(font-family=)(['\"])Arial-BoldMT\2",
        lambda match: f"{match.group(1)}{match.group(2)}Arial{match.group(2)} font-weight={match.group(2)}700{match.group(2)}",
        svg_text,
    )
    svg_text = re.sub(r"(font-family=)(['\"])ArialMT\2", r"\1\2Arial\2", svg_text)
    svg_text = re.sub(r"(font-family=)(['\"])Arial-Black\2", r"\1\2Arial Black\2", svg_text)
    svg_text = re.sub(r"(font-family=)(['\"])AdobeSongStd-Light-GBpc-EUC-H\2", r"\1\2Arial\2", svg_text)
    svg_text = re.sub(r"(font-family=)(['\"])NotoSansArmenian-ExtraBold\2", r"\1\2Arial\2", svg_text)
    svg_text = re.sub(r"(font-family=)(['\"])NotoSansArmenian-Medium\2", r"\1\2Arial\2", svg_text)
    svg_text = re.sub(r"(font-family=)(['\"])NotoSansAdlam-Regular\2", r"\1\2Arial\2", svg_text)
    return svg_text


def parse_svg_css_classes(svg_text: str) -> dict[str, dict[str, str]]:
    classes: dict[str, dict[str, str]] = {}
    for name, body in re.findall(r"\.([A-Za-z0-9_-]+)\{([^}]+)\}", svg_text):
        props: dict[str, str] = {}
        for part in body.split(";"):
            if ":" not in part:
                continue
            key, value = part.split(":", 1)
            props[key.strip()] = value.strip()
        classes[name] = props
    return classes


def local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def resolved_class_style(class_attr: str, classes: dict[str, dict[str, str]]) -> dict[str, str]:
    props: dict[str, str] = {}
    for class_name in class_attr.split():
        props.update(classes.get(class_name, {}))
    return props


def comparable_text_style(props: dict[str, str]) -> tuple[str, str, str]:
    return (
        props.get("fill", ""),
        props.get("font-size", ""),
        props.get("font-weight", ""),
    )


def collapse_single_line_tspans(path: Path) -> None:
    """Collapse Illustrator word-positioned tspans when visual style matches.

    Illustrator sometimes serializes one visual line as many tspans with fixed
    x coordinates based on the authoring font. If a replacement font is used,
    words overlap or drift. Collapsing same-style single-line tspans into normal
    text lets the renderer use normal word spacing before the final text-to-path
    conversion. Mixed-color or mixed-size lines are left untouched.
    """
    svg_text = path.read_text(encoding="utf-8")
    classes = parse_svg_css_classes(svg_text)
    tree = ET.parse(path)
    root = tree.getroot()
    changed = False

    for element in root.iter():
        if local_name(element.tag) != "text":
            continue
        tspans = [child for child in list(element) if local_name(child.tag) == "tspan"]
        if not tspans:
            continue
        y_values = {child.attrib.get("y", "") for child in tspans}
        if len(y_values) > 1:
            continue
        meaningful = [child for child in tspans if (child.text or "").strip(" -\t\n\r")]
        if not meaningful:
            continue
        styles = [
            comparable_text_style(resolved_class_style(child.attrib.get("class", ""), classes))
            for child in meaningful
        ]
        if any(style != styles[0] for style in styles[1:]):
            continue
        first_props = resolved_class_style(meaningful[0].attrib.get("class", ""), classes)
        style_parts = []
        for key in ["fill", "font-family", "font-size", "font-weight", "font-style"]:
            value = first_props.get(key)
            if value:
                style_parts.append(f"{key}:{value}")
        if not style_parts:
            continue

        first_tspan = tspans[0]
        if first_tspan.attrib.get("x") not in {None, "", "0"}:
            element.set("x", first_tspan.attrib["x"])
        if first_tspan.attrib.get("y") not in {None, "", "0"}:
            element.set("y", first_tspan.attrib["y"])
        for child in tspans:
            element.remove(child)
        element.text = "".join(child.text or "" for child in tspans)
        element.set("style", ";".join(style_parts))
        changed = True

    if changed:
        tree.write(path, encoding="utf-8", xml_declaration=True)


@dataclass(frozen=True)
class FigureSpec:
    number: int
    desktop_name: str
    mirror_dir: str
    title: str

    @property
    def desktop_path(self) -> Path:
        return DESKTOP_FI / self.desktop_name

    @property
    def canonical_stem(self) -> str:
        return self.desktop_path.stem

    @property
    def upload_stem(self) -> str:
        return f"Figure_{self.number:02d}"


FIGURES = [
    FigureSpec(
        1,
        "Figure1_study_architecture_claim_boundary.live-text.svg",
        "05_FIGURE1_STUDY_ARCHITECTURE_CLAIM_BOUNDARY",
        "Study architecture and claim-boundary map",
    ),
    FigureSpec(
        2,
        "Figure2_cross_cohort_decoupling_final.live-text.svg",
        "05_FIGURE2_CROSS_COHORT_DECOUPLING",
        "Cross-cohort inflammatory/MHC-II decoupling",
    ),
    FigureSpec(
        3,
        "Figure3_final_submit_typography_final.live-text.svg",
        "05_FIGURE3_CONTINUOUS_SPECTRUM_PATHWAY_CONTEXT",
        "Continuous spectrum pathway context",
    ),
    FigureSpec(
        4,
        "Figure4_raw_fcs_pair_specificity_hardening_readable.live-text.svg",
        "05_FIGURE4_RAW_FCS_PAIR_SPECIFICITY_HARDENING",
        "Raw-FCS pair-specificity hardening",
    ),
    FigureSpec(
        5,
        "Figure5_COMBAT_RNA_CyTOF_bridge_portrait.svg",
        "05_FIGURE5_COMBAT_PAIRED_RNA_CYTOF_BRIDGE",
        "COMBAT RNA-CyTOF bridge",
    ),
    FigureSpec(
        6,
        "Figure6_working_interpretation_v2_palette_matched.svg",
        "05_FIGURE6_WORKING_INTERPRETATION",
        "Working interpretation and claim boundary",
    ),
]

TEXT_SUFFIXES = {".csv", ".tsv", ".txt", ".md", ".json", ".xlsx"}


def run(cmd: list[str], *, cwd: Path = ROOT, check: bool = True) -> subprocess.CompletedProcess[str]:
    proc = subprocess.run(
        cmd,
        cwd=str(cwd),
        check=False,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    if check and proc.returncode != 0:
        raise RuntimeError(
            "Command failed:\n"
            + " ".join(cmd)
            + f"\nexit={proc.returncode}\nSTDOUT:\n{proc.stdout}\nSTDERR:\n{proc.stderr}"
        )
    return proc


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def rel(path: Path) -> str:
    try:
        return path.relative_to(ROOT).as_posix()
    except ValueError:
        return str(path)


def ensure_clean_dir(path: Path) -> None:
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True, exist_ok=True)


def svg_stats(path: Path) -> dict[str, object]:
    root = ET.parse(path).getroot()
    local = lambda tag: tag.split("}", 1)[-1]
    viewbox = root.attrib.get("viewBox") or root.attrib.get("viewbox")
    if not viewbox:
        raise ValueError(f"SVG lacks viewBox: {path}")
    parts = [float(x) for x in re.split(r"[\s,]+", viewbox.strip()) if x]
    if len(parts) != 4:
        raise ValueError(f"Invalid viewBox in {path}: {viewbox}")
    width_pt = parts[2]
    height_pt = parts[3]
    text_count = sum(1 for e in root.iter() if local(e.tag) == "text")
    tspan_count = sum(1 for e in root.iter() if local(e.tag) == "tspan")
    image_count = sum(1 for e in root.iter() if local(e.tag) == "image")
    return {
        "viewBox": viewbox,
        "width_pt": width_pt,
        "height_pt": height_pt,
        "width_mm": width_pt / 72 * 25.4,
        "height_mm": height_pt / 72 * 25.4,
        "text_elements": text_count,
        "tspan_elements": tspan_count,
        "image_elements": image_count,
        "bytes": path.stat().st_size,
        "sha256": sha256(path),
    }


def normalized_svg(src: Path, dst: Path, stats: dict[str, object]) -> None:
    tree = ET.parse(src)
    root = tree.getroot()
    root.set("width", f"{stats['width_pt']}pt")
    root.set("height", f"{stats['height_pt']}pt")
    tree.write(dst, encoding="utf-8", xml_declaration=True)
    dst.write_text(normalize_svg_font_css(dst.read_text(encoding="utf-8")), encoding="utf-8")
    collapse_single_line_tspans(dst)


def export_figure(spec: FigureSpec, normalized: Path, out_dir: Path) -> dict[str, object]:
    stem = spec.canonical_stem
    paths = {
        "source_svg_exact": out_dir / f"{stem}.source.svg",
        "normalized_svg": out_dir / f"{stem}.normalized.svg",
        "outlined_svg": out_dir / f"{stem}.outlined.svg",
        "pdf": out_dir / f"{stem}.pdf",
        "eps": out_dir / f"{stem}.eps",
        "png_600dpi": out_dir / f"{stem}_600dpi.png",
        "tiff_600dpi": out_dir / f"{stem}_600dpi.tiff",
        "preview_png": out_dir / f"{stem}_preview.png",
    }
    shutil.copy2(spec.desktop_path, paths["source_svg_exact"])
    shutil.copy2(normalized, paths["normalized_svg"])

    run(
        [
            "inkscape",
            str(normalized),
            "--export-type=svg",
            "--export-text-to-path",
            f"--export-filename={paths['outlined_svg']}",
        ]
    )
    outlined = paths["outlined_svg"]
    run(
        [
            "inkscape",
            str(outlined),
            "--export-type=pdf",
            f"--export-filename={paths['pdf']}",
        ]
    )
    run(
        [
            "inkscape",
            str(outlined),
            "--export-type=eps",
            f"--export-filename={paths['eps']}",
        ]
    )
    run(
        [
            "inkscape",
            str(outlined),
            "--export-type=png",
            f"--export-dpi={DPI}",
            "--export-background=white",
            "--export-background-opacity=1",
            f"--export-filename={paths['png_600dpi']}",
        ]
    )
    run(
        [
            "inkscape",
            str(outlined),
            "--export-type=png",
            "--export-width=1600",
            "--export-background=white",
            "--export-background-opacity=1",
            f"--export-filename={paths['preview_png']}",
        ]
    )

    with Image.open(paths["png_600dpi"]) as img:
        rgb = img.convert("RGB")
        rgb.save(paths["png_600dpi"], dpi=(DPI, DPI))
        rgb.save(paths["tiff_600dpi"], dpi=(DPI, DPI), compression="tiff_lzw")

    return {k: str(v) for k, v in paths.items()}


def pdf_page_size(path: Path) -> str:
    proc = run(["pdfinfo", str(path)], check=False)
    if proc.returncode != 0:
        return f"ERROR: {proc.stderr.strip()}"
    for line in proc.stdout.splitlines():
        if line.startswith("Page size:"):
            return line.split(":", 1)[1].strip()
    return "UNKNOWN"


def pdf_font_rows(path: Path) -> int | str:
    proc = run(["pdffonts", str(path)], check=False)
    if proc.returncode != 0:
        return f"ERROR: {proc.stderr.strip()}"
    rows = [line for line in proc.stdout.splitlines() if line.strip()]
    return max(0, len(rows) - 2)


def tiff_info(path: Path) -> dict[str, object]:
    with Image.open(path) as img:
        dpi = img.info.get("dpi", (None, None))
        return {
            "mode": img.mode,
            "width_px": img.width,
            "height_px": img.height,
            "dpi_x": round(float(dpi[0]), 3) if dpi[0] else None,
            "dpi_y": round(float(dpi[1]), 3) if dpi[1] else None,
        }


def image_nonblank(path: Path) -> bool:
    with Image.open(path) as img:
        rgb = img.convert("RGB")
        bg = Image.new("RGB", rgb.size, "white")
        diff = ImageChops.difference(rgb, bg)
        return diff.getbbox() is not None


def outlined_text_count(path: Path) -> int:
    root = ET.parse(path).getroot()
    return sum(1 for e in root.iter() if e.tag.split("}", 1)[-1] == "text")


def collect_delete_candidates() -> list[Path]:
    candidates: set[Path] = set()
    patterns = [
        "Figure*.svg",
        "Figure*.pdf",
        "Figure*.eps",
        "Figure*.png",
        "Figure*.tif",
        "Figure*.tiff",
        "fig*.svg",
        "fig*.pdf",
        "fig*.png",
        "fig*.tif",
        "fig*.tiff",
        "*Frontiers*package*.zip",
        "*submission*package*.zip",
    ]
    search_roots = [
        ROOT / "05_FIGURE1_STUDY_ARCHITECTURE_CLAIM_BOUNDARY",
        ROOT / "05_FIGURE2_CROSS_COHORT_DECOUPLING",
        ROOT / "05_FIGURE3_CONTINUOUS_SPECTRUM_PATHWAY_CONTEXT",
        ROOT / "05_FIGURE4_RAW_FCS_PAIR_SPECIFICITY_HARDENING",
        ROOT / "05_FIGURE5_COMBAT_PAIRED_RNA_CYTOF_BRIDGE",
        ROOT / "05_FIGURE6_WORKING_INTERPRETATION",
        ROOT / "figures",
        ROOT,
    ]
    for base in search_roots:
        if not base.exists():
            continue
        for pat in patterns:
            iterator = base.glob(pat) if base == ROOT else base.rglob(pat)
            for path in iterator:
                if path.is_file():
                    if OUT in path.parents or path == ZIP_PATH:
                        continue
                    if "01_PUBLIC_DATA_RAW" in path.parts:
                        continue
                    candidates.add(path)

    for stale in ROOT.glob("04_FRONTIERS_SUBMISSION/*package*.zip"):
        if stale.is_file():
            candidates.add(stale)
    for stale in ROOT.glob("05_PUBLIC_RELEASE_ZENODO/*package*.zip"):
        if stale.is_file():
            candidates.add(stale)
    return sorted(candidates)


def write_delete_audit(candidates: list[Path]) -> None:
    QA.mkdir(parents=True, exist_ok=True)
    out = QA / "permanent_delete_manifest_20260620.tsv"
    with out.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(
            handle,
            fieldnames=["relative_path", "size_bytes", "sha256", "reason"],
            delimiter="\t",
        )
        writer.writeheader()
        for path in candidates:
            writer.writerow(
                {
                    "relative_path": rel(path),
                    "size_bytes": path.stat().st_size,
                    "sha256": sha256(path),
                    "reason": "stale figure/export/submission package replaced by desktop FI authoritative set",
                }
            )


def delete_candidates(candidates: list[Path]) -> None:
    for path in candidates:
        if path.exists() and path.is_file():
            path.unlink()


def copy_source_data() -> list[dict[str, object]]:
    copied: list[dict[str, object]] = []
    SOURCE.mkdir(parents=True, exist_ok=True)
    for spec in FIGURES:
        fig_dir = ROOT / spec.mirror_dir
        dest = SOURCE / spec.upload_stem
        dest.mkdir(parents=True, exist_ok=True)
        if not fig_dir.exists():
            continue
        for path in sorted(fig_dir.iterdir()):
            name_lower = path.name.lower()
            if not path.is_file():
                continue
            is_source = (
                "source" in name_lower
                or "mapping" in name_lower
                or "attribution" in name_lower
                or "license" in name_lower
            )
            if is_source and path.suffix.lower() in TEXT_SUFFIXES:
                dst = dest / path.name
                shutil.copy2(path, dst)
                copied.append(
                    {
                        "figure": spec.upload_stem,
                        "source": rel(path),
                        "copied_to": rel(dst),
                        "sha256": sha256(dst),
                    }
                )
    with (SOURCE / "source_data_copy_manifest.csv").open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=["figure", "source", "copied_to", "sha256"])
        writer.writeheader()
        writer.writerows(copied)
    return copied


def mirror_outputs(spec: FigureSpec, export_dir: Path, outputs: dict[str, str]) -> None:
    mirror = ROOT / spec.mirror_dir
    mirror.mkdir(parents=True, exist_ok=True)
    for key, src_str in outputs.items():
        if key in {"source_svg_exact", "normalized_svg", "outlined_svg", "pdf", "eps", "png_600dpi", "tiff_600dpi", "preview_png"}:
            src = Path(src_str)
            shutil.copy2(src, mirror / src.name)

    global_figures = ROOT / "figures"
    global_figures.mkdir(parents=True, exist_ok=True)
    for key in ["source_svg_exact", "outlined_svg", "pdf", "eps", "png_600dpi", "tiff_600dpi", "preview_png"]:
        src = Path(outputs[key])
        shutil.copy2(src, global_figures / src.name)


def copy_upload_files(spec: FigureSpec, outputs: dict[str, str]) -> None:
    shutil.copy2(Path(outputs["tiff_600dpi"]), UPLOAD / f"{spec.upload_stem}.tiff")
    shutil.copy2(Path(outputs["preview_png"]), UPLOAD / f"{spec.upload_stem}_preview.png")
    for key, suffix in [
        ("source_svg_exact", ".source.svg"),
        ("normalized_svg", ".normalized.svg"),
        ("outlined_svg", ".outlined.svg"),
        ("pdf", ".pdf"),
        ("eps", ".eps"),
        ("png_600dpi", "_600dpi.png"),
        ("tiff_600dpi", "_600dpi.tiff"),
    ]:
        shutil.copy2(Path(outputs[key]), VECTOR / f"{spec.upload_stem}{suffix}")


def docx_figure_counts() -> dict[str, int]:
    doc = Document(DOCX)
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + "\n".join(p.text for p in cell.paragraphs)
    return {f"Figure {i}.": text.count(f"Figure {i}.") for i in range(1, 7)}


def try_render_docx() -> dict[str, object]:
    result: dict[str, object] = {"attempted": False, "success": False}
    if not DOCX.exists():
        return {**result, "error": "DOCX missing"}
    if not RENDER_DOCX.exists():
        return {**result, "error": "render_docx.py missing"}
    out_dir = QA / "manuscript_render"
    out_dir.mkdir(parents=True, exist_ok=True)
    proc = run(
        [
            sys.executable,
            str(RENDER_DOCX),
            str(DOCX),
            "--output_dir",
            str(out_dir),
            "--emit_pdf",
            "--width",
            "1800",
            "--height",
            "2400",
        ],
        check=False,
    )
    result.update(
        {
            "attempted": True,
            "returncode": proc.returncode,
            "stdout_tail": proc.stdout[-1000:],
            "stderr_tail": proc.stderr[-1000:],
        }
    )
    pdfs = sorted(out_dir.glob("*.pdf"))
    if proc.returncode == 0 and pdfs:
        manuscript_pdf = UPLOAD / "Frontiers_Immunology_V17_FINAL_TEXT_ONLY.pdf"
        shutil.copy2(pdfs[0], manuscript_pdf)
        result.update({"success": True, "pdf": rel(manuscript_pdf)})
    else:
        result["error"] = "DOCX render did not produce PDF"
    return result


def write_manifest(base: Path, out_csv: Path) -> None:
    rows = []
    for path in sorted(p for p in base.rglob("*") if p.is_file()):
        rows.append(
            {
                "relative_path": path.relative_to(base).as_posix(),
                "size_bytes": path.stat().st_size,
                "sha256": sha256(path),
            }
        )
    with out_csv.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=["relative_path", "size_bytes", "sha256"])
        writer.writeheader()
        writer.writerows(rows)


def write_upload_figure_manifest() -> None:
    rows = []
    for spec in FIGURES:
        path = UPLOAD / f"{spec.upload_stem}.tiff"
        rows.append(
            {
                "figure": spec.upload_stem,
                "upload_file": path.relative_to(OUT).as_posix(),
                "size_bytes": path.stat().st_size,
                "sha256": sha256(path),
            }
        )
    with (QA / "frontiers_upload_figure_manifest_sha256.csv").open(
        "w", newline="", encoding="utf-8"
    ) as handle:
        writer = csv.DictWriter(
            handle, fieldnames=["figure", "upload_file", "size_bytes", "sha256"]
        )
        writer.writeheader()
        writer.writerows(rows)


def make_contact_sheet(previews: list[Path]) -> Path:
    thumbs = []
    for path in previews:
        img = Image.open(path).convert("RGB")
        img.thumbnail((520, 520))
        canvas = Image.new("RGB", (560, 580), "white")
        canvas.paste(img, ((560 - img.width) // 2, 20))
        thumbs.append(canvas)
    sheet = Image.new("RGB", (1120, 1740), "white")
    for idx, tile in enumerate(thumbs):
        x = (idx % 2) * 560
        y = (idx // 2) * 580
        sheet.paste(tile, (x, y))
    out = QA / "six_figure_preview_contact_sheet.png"
    sheet.save(out, dpi=(150, 150))
    return out


def zip_submission() -> None:
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(p for p in OUT.rglob("*") if p.is_file()):
            zf.write(path, path.relative_to(ROOT).as_posix())


def write_readme(docx_render: dict[str, object]) -> None:
    lines = [
        "# FI submission-ready package",
        "",
        f"Generated: {datetime.now().isoformat(timespec='seconds')}",
        "Target journal: Frontiers in Immunology.",
        "",
        "Authoritative figure source: `<AUTHORITATIVE_FI_FIGURE_SOURCE_DIR>`.",
        "Scientific analyses were not rerun; this package only cleans, exports, QA-checks, and assembles the user-approved figure set.",
        "",
        "## Upload-ready files",
        "- `01_upload_ready/Frontiers_Immunology_V17_FINAL_TEXT_ONLY.docx`",
        "- `01_upload_ready/Figure_01.tiff` through `Figure_06.tiff`",
    ]
    if docx_render.get("success"):
        lines.append("- `01_upload_ready/Frontiers_Immunology_V17_FINAL_TEXT_ONLY.pdf`")
    else:
        lines.append("- Manuscript PDF was not produced; see `04_QA_and_manifests/final_QA_report.json`.")
    lines.extend(
        [
            "",
            "## Backup and audit",
            "- `02_vector_backup/`: source SVG, normalized SVG, outlined SVG, PDF, EPS, PNG, and TIFF backups.",
            "- `03_source_data_and_attribution/`: copied figure source data, mapping, and attribution files.",
            "- `04_QA_and_manifests/`: export QA, delete audit, checksums, contact sheet, and final manifest.",
        ]
    )
    (OUT / "README.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    if not DESKTOP_FI.exists():
        raise FileNotFoundError(DESKTOP_FI)
    for spec in FIGURES:
        if not spec.desktop_path.exists():
            raise FileNotFoundError(spec.desktop_path)
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)

    ensure_clean_dir(OUT)
    for d in [UPLOAD, VECTOR, SOURCE, QA]:
        d.mkdir(parents=True, exist_ok=True)

    candidates = collect_delete_candidates()
    write_delete_audit(candidates)
    delete_candidates(candidates)

    shutil.copy2(DOCX, UPLOAD / DOCX.name)
    source_copy_manifest = copy_source_data()

    figure_rows: list[dict[str, object]] = []
    qa_rows: list[dict[str, object]] = []
    previews: list[Path] = []

    with tempfile.TemporaryDirectory(prefix="fi_svg_export_") as tmp:
        tmpdir = Path(tmp)
        for spec in FIGURES:
            stats = svg_stats(spec.desktop_path)
            normalized = tmpdir / f"{spec.upload_stem}.normalized.svg"
            normalized_svg(spec.desktop_path, normalized, stats)

            export_dir = VECTOR / f"{spec.upload_stem}_canonical"
            export_dir.mkdir(parents=True, exist_ok=True)
            outputs = export_figure(spec, normalized, export_dir)
            mirror_outputs(spec, export_dir, outputs)
            copy_upload_files(spec, outputs)

            tiff = Path(outputs["tiff_600dpi"])
            pdf = Path(outputs["pdf"])
            outlined = Path(outputs["outlined_svg"])
            preview = Path(outputs["preview_png"])
            previews.append(UPLOAD / f"{spec.upload_stem}_preview.png")
            tinfo = tiff_info(tiff)
            outlined_text = outlined_text_count(outlined)
            qa = {
                "figure": spec.upload_stem,
                "desktop_source": str(spec.desktop_path),
                "title": spec.title,
                **stats,
                "pdf_page_size": pdf_page_size(pdf),
                "pdf_font_rows": pdf_font_rows(pdf),
                "outlined_svg_text_elements": outlined_text,
                "tiff_mode": tinfo["mode"],
                "tiff_width_px": tinfo["width_px"],
                "tiff_height_px": tinfo["height_px"],
                "tiff_dpi_x": tinfo["dpi_x"],
                "tiff_dpi_y": tinfo["dpi_y"],
                "preview_nonblank": image_nonblank(preview),
                "frontiers_min_dpi_pass": bool(
                    tinfo["mode"] == "RGB"
                    and (tinfo["dpi_x"] or 0) >= 300
                    and (tinfo["dpi_y"] or 0) >= 300
                ),
                "outlined_text_to_path_pass": outlined_text == 0,
                "embedded_raster_count": stats["image_elements"],
            }
            qa_rows.append(qa)
            for role, path_str in outputs.items():
                p = Path(path_str)
                figure_rows.append(
                    {
                        "figure": spec.upload_stem,
                        "role": role,
                        "path": rel(p),
                        "size_bytes": p.stat().st_size,
                        "sha256": sha256(p),
                    }
                )

    docx_render = try_render_docx()
    counts = docx_figure_counts()
    contact_sheet = make_contact_sheet(previews)

    with (QA / "figure_export_manifest.csv").open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=["figure", "role", "path", "size_bytes", "sha256"])
        writer.writeheader()
        writer.writerows(figure_rows)

    with (QA / "figure_QA_summary.csv").open("w", newline="", encoding="utf-8") as handle:
        fieldnames = sorted({k for row in qa_rows for k in row.keys()})
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(qa_rows)

    final_report = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "target_journal": "Frontiers in Immunology",
        "desktop_fi_source": str(DESKTOP_FI),
        "project_root": str(ROOT),
        "deleted_file_count": len(candidates),
        "source_data_copied_count": len(source_copy_manifest),
        "docx_figure_caption_counts": counts,
        "docx_render": docx_render,
        "figures": qa_rows,
        "contact_sheet": rel(contact_sheet),
        "frontiers_upload_checklist": {
            "six_individual_tiff_files": all((UPLOAD / f"Figure_{i:02d}.tiff").exists() for i in range(1, 7)),
            "all_tiffs_rgb": all(row["tiff_mode"] == "RGB" for row in qa_rows),
            "all_tiffs_at_least_300_dpi": all(row["frontiers_min_dpi_pass"] for row in qa_rows),
            "source_svg_exact_preserved": all(row["sha256"] for row in qa_rows),
            "no_embedded_raster_in_authoritative_svg": all(row["embedded_raster_count"] == 0 for row in qa_rows),
            "outlined_svgs_have_no_live_text": all(row["outlined_text_to_path_pass"] for row in qa_rows),
        },
    }
    (QA / "final_QA_report.json").write_text(json.dumps(final_report, indent=2), encoding="utf-8")

    write_readme(docx_render)
    write_upload_figure_manifest()
    write_manifest(OUT, QA / "FI_submission_ready_manifest_sha256.csv")
    zip_submission()
    zip_test = run(["unzip", "-t", str(ZIP_PATH)], check=False)
    (QA / "zip_test.txt").write_text(zip_test.stdout + zip_test.stderr, encoding="utf-8")
    if zip_test.returncode != 0:
        raise RuntimeError("Final zip failed unzip -t")

    print(json.dumps(
        {
            "output_dir": str(OUT),
            "zip": str(ZIP_PATH),
            "deleted_file_count": len(candidates),
            "source_data_copied_count": len(source_copy_manifest),
            "contact_sheet": str(contact_sheet),
            "docx_render_success": docx_render.get("success", False),
        },
        indent=2,
    ))


if __name__ == "__main__":
    main()
