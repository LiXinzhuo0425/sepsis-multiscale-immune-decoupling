#!/usr/bin/env python3
"""Final light language polish for the revised FI manuscript.

Scope is intentionally narrow:
- Contribution to the Field
- Abstract
- final Introduction paragraph
- first six Discussion paragraphs
- Figure 1, 4, 5 and 6 legends
- Generative AI statement

The pass reduces repeated defensive phrasing without changing methods, main
results, quantitative values or claim boundaries.
"""

from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "07_FI_SUBMISSION_REVISED_20260622"
UPLOAD = OUT / "01_upload_ready"
QA = OUT / "04_QA_and_manifests"
MAIN_DOCX = ROOT / "00_FINAL_WORD/Frontiers_Immunology_V21_FI_SUBMISSION_MAIN_TEXT_ONLY_REVISED.docx"
UPLOAD_MAIN_DOCX = UPLOAD / MAIN_DOCX.name
ZIP_PATH = ROOT / "FI_submission_revised_20260622.zip"
DOCX_RENDERER = Path(
    "<CODEX_DOCUMENTS_RENDER_DOCX>"
)
SOFFICE_APP = Path("/Applications/LibreOffice.app/Contents/MacOS")

LONG_PHRASE = "representative residual co-event signal within a broader marker-pair co-event background"

REPLACEMENTS = {
    # Contribution to the Field
    (
        "Sepsis often combines inflammatory activation with impaired antigen presentation, but RNA profiling and cytometry measure different immune layers. "
        "This study asks how those layers relate using public datasets rather than new samples. Across seven public sepsis RNA cohorts, inflammatory programs increased while antigen-presentation programs involving MHC-II, CD74 and HLA-DR decreased or inversely coupled with inflammation. "
        "In COMBAT, paired RNA and mass cytometry measurements linked this RNA pattern to a representative residual co-event signal within a broader marker-pair co-event background. "
        "Raw-FCS analyses showed that the prespecified CD3/CD14 metric was not uniquely specific, so the cytometry result is interpreted as an artifact-aware event-derived summary rather than a structural interaction or new cell subset. "
        "Case-only analyses supported a continuous immune-decoupling gradient, and donor-level single-cell and pathway analyses supplied context. The contribution is a reproducible, claim-bounded framework for integrating public RNA, cytometry and single-cell layers in sepsis. "
        "By separating recurrence, artifact controls and claim boundaries, it clarifies what public multi-layer data can and cannot support. It is association-only and requires prospective paired RNA, cytometry and protein-level validation before clinical or mechanistic use."
    ): (
        "Sepsis often combines inflammatory activation with impaired antigen presentation, but RNA profiling and cytometry measure different immune layers. "
        "This study used public datasets to ask how those layers relate. Across seven sepsis RNA cohorts, inflammatory programs increased while antigen-presentation programs involving MHC-II, CD74 and HLA-DR decreased or inversely coupled with inflammation. "
        "In COMBAT, paired RNA and mass cytometry linked this RNA pattern to a residual co-event readout. "
        "Raw-FCS analyses showed that the prespecified CD3/CD14 metric was not uniquely specific, so the cytometry result is treated as an event-derived summary rather than a structural interaction or new cell subset. "
        "Case-only analyses supported a continuous immune-decoupling gradient, and donor-level single-cell and pathway analyses supplied context. The contribution is a reproducible integration strategy for public RNA, cytometry and single-cell layers in sepsis. "
        "By separating cross-cohort recurrence, artifact controls and biological interpretation, the study clarifies what public multi-layer data can and cannot support. "
        "The findings remain association-only and require prospective paired RNA, cytometry and protein-level validation before clinical or mechanistic use."
    ),
    # Abstract
    (
        "Background: Whole-blood RNA sequencing and cytometry assays capture different immune analytical layers, which can obscure how inflammation and immunoparalysis coexist in sepsis. "
        "Methods: We integrated seven public bulk sepsis transcriptomic cohorts, derived cytometry summary tables, COMBAT paired RNA sequencing and mass cytometry participant-timepoint data, matched raw-FCS sensitivity analyses, and donor-level single-cell contextual analysis. "
        "Results: A total of 660 bulk RNA samples (499 sepsis cases) were included. The prespecified six-gene host-response panel showed inverse coupling with MHC-II/CD74 across all seven cohorts (median Spearman rho approximately -0.60), and composition sensitivity attenuated but did not reverse the primary direction. "
        "Case-only analysis supported a continuous immune-decoupling gradient rather than a validated clinical endotype. In COMBAT sepsis-only paired samples (40 rows from 34 participants), a representative residual co-event signal within a broader marker-pair co-event background, operationalized by the prespecified CD3/CD14 metric, was negatively associated with HLA-DR core RNA score (rho = -0.559, permutation p = 4.00 x 10^-4) and MHC-II/CD74 RNA score (rho = -0.517, permutation p = 0.0011), and positively associated with RNA decoupling index (rho = 0.462, permutation p = 0.0035). "
        "Row-level, participant-level, participant-cluster, and composition-adjusted sensitivities preserved association directions with expected attenuation. Raw-FCS all-pair analysis did not resolve marker-pair specificity; therefore the cytometry finding is interpreted as representative and artifact-aware rather than CD3/CD14-specific. "
        "Conclusion: These findings support a bounded systems immunology reconstruction of sepsis-associated immunoparalysis that requires prospective paired biological validation."
    ): (
        "Background: Whole-blood RNA sequencing and cytometry assays capture different immune analytical layers, which can obscure how inflammation and immunoparalysis coexist in sepsis. "
        "Methods: We integrated seven public bulk sepsis transcriptomic cohorts, derived cytometry summary tables, COMBAT paired RNA sequencing and mass cytometry participant-timepoint data, matched raw-FCS sensitivity analyses, and donor-level single-cell contextual analysis. "
        "Results: A total of 660 bulk RNA samples (499 sepsis cases) were included. The prespecified six-gene host-response panel showed inverse coupling with MHC-II/CD74 across all seven cohorts (median Spearman rho approximately -0.60), and composition sensitivity attenuated but did not reverse the primary direction. "
        "Case-only analysis supported a continuous immune-decoupling gradient rather than a validated clinical endotype. In COMBAT sepsis-only paired samples (40 rows from 34 participants), a representative residual co-event signal within a broader marker-pair co-event background, measured by the prespecified CD3/CD14 metric, was negatively associated with HLA-DR core RNA score (rho = -0.559, permutation p = 4.00 x 10^-4) and MHC-II/CD74 RNA score (rho = -0.517, permutation p = 0.0011), and positively associated with RNA decoupling index (rho = 0.462, permutation p = 0.0035). "
        "Row-level, participant-level, participant-cluster, and composition-adjusted sensitivities preserved association directions with expected attenuation. Raw-FCS all-pair analysis did not resolve marker-pair specificity, so this cytometry readout is interpreted as representative and artifact-aware rather than CD3/CD14-specific. "
        "Conclusion: These findings support a conservative systems-level interpretation of sepsis-associated immunoparalysis that requires prospective paired biological validation."
    ),
    # Introduction final paragraph
    (
        "Here we investigated whether public RNA and cytometry data could support a bounded, compartment-aware model of sepsis-associated immunoparalysis. "
        "The study architecture and claim boundaries are summarized in Figure 1. Our evidence hierarchy proceeds from cross-cohort bulk RNA recurrence to a continuous case-only immune-decoupling spectrum, raw-FCS marker-pair sensitivity around a representative residual co-event signal, a paired COMBAT RNA-CyTOF bridge, donor-level single-cell contextual support, pathway context and an exploratory 28-day mortality clinical anchor. "
        "The central conclusion is constrained to the following: public RNA and cytometry data support a compartment-aware model of sepsis-associated immunoparalysis, in which systemic inflammatory/MHC-II decoupling is associated with a representative residual co-event signal within a broader marker-pair co-event background, while the event-derived signal itself remains an artifact-aware summary rather than a specific structural interaction or validated cell subset."
    ): (
        "Here we investigated whether public RNA and cytometry data could support a compartment-aware interpretation of sepsis-associated immunoparalysis. "
        "The study architecture and interpretive limits are summarized in Figure 1. Our evidence hierarchy proceeds from cross-cohort bulk RNA recurrence to a continuous case-only immune-decoupling spectrum, raw-FCS marker-pair sensitivity around a representative residual co-event signal, a paired COMBAT RNA-CyTOF bridge, donor-level single-cell context, pathway context and an exploratory 28-day mortality clinical anchor. "
        "We therefore interpret the evidence conservatively: public RNA and cytometry data support a compartment-aware model in which systemic inflammatory/MHC-II decoupling is associated with the cytometry readout in paired sepsis samples. The event-derived signal is treated as a summary measure, not as a specific structural interaction or validated cell subset."
    ),
    # Discussion first six paragraphs
    (
        "This study supports a bounded systems immunology reconstruction of sepsis-associated immunoparalysis. "
        "The fundamental contribution is not the demonstration that MHC-II or HLA-DR suppression exists in sepsis, as this is already well established, but rather that systemic inflammatory/MHC-II decoupling can be associated with a representative residual co-event signal within a broader marker-pair co-event background in matched sepsis samples, while the residual event-derived summary itself exhibits activation-state context."
    ): (
        "This study supports a conservative systems-level interpretation of sepsis-associated immunoparalysis. "
        "The contribution is not to show that MHC-II or HLA-DR suppression exists in sepsis, as this is already well established. Instead, the matched public data suggest that systemic inflammatory/MHC-II decoupling can track with a residual co-event readout, while the event-derived summary itself carries activation-state context."
    ),
    (
        "The paired COMBAT analytical layer remains the key distinction from a standard public transcriptomic study, and the expanded sensitivity analyses further constrain its interpretation. "
        "It provides a matched public-data bridge between whole-blood RNA and CyTOF summaries within participant-timepoints. This bridge supports the proposed systems immunology reconstruction but cannot serve as clinical validation or establish causal relationships. "
        "Composition-adjusted, participant-level and lineage/QC-residualized sensitivities show that the association direction is not a trivial artifact of one analysis scale, while the attenuation after stronger adjustment argues for conservative interpretation."
    ): (
        "The paired COMBAT layer is the main distinction from a standard public transcriptomic study, and the expanded sensitivity analyses define how far it can be interpreted. "
        "It provides a matched public-data bridge between whole-blood RNA and CyTOF summaries within participant-timepoints. The bridge supports the proposed interpretation but cannot serve as clinical validation or establish causal relationships. "
        "Composition-adjusted, participant-level and lineage/QC-residualized sensitivities argue against a single-scale artifact, while attenuation after stronger adjustment keeps the interpretation cautious."
    ),
    (
        "The cytometry evidence also necessitates conservative reformulation. Raw-FCS pair-specificity hardening did not establish CD3/CD14 as a uniquely specific pair; rather, CD3/CD14 remains a representative residual co-event signal within a broader marker-pair co-event background. "
        "The safer interpretation is one of analytical-layer divergence: whole-blood RNA can capture population-level antigen-presentation suppression while residual event-derived summaries show summary-level marker-state context and co-event frequency behavior that are not reducible to a single structural interaction claim."
    ): (
        "The cytometry evidence also requires a conservative interpretation. Raw-FCS pair-specificity analyses did not establish CD3/CD14 as a uniquely specific pair. "
        "CD3/CD14 is retained as a representative readout within a wider marker-pair background. The key point is analytical-layer divergence: whole-blood RNA can capture population-level antigen-presentation suppression while residual event-derived summaries show marker-state context and co-event frequency behavior that are not reducible to a single structural interaction claim."
    ),
    (
        "The stronger CD14/CD16 and CD4/CD14 associations provide a biologically plausible caution against over-centering a single pair: CD14/CD16 may capture broader myeloid and monocyte-lineage co-event background, whereas CD4/CD14 may be closer to helper T-cell/monocyte antigen-presentation context. "
        "Accordingly, the study treats CD3/CD14 as a representative residual co-event signal within a broader marker-pair co-event background rather than as a uniquely specific pair."
    ): (
        "The stronger CD14/CD16 and CD4/CD14 associations provide a biologically plausible caution against over-centering one marker pair. "
        "CD14/CD16 may capture broader myeloid and monocyte-lineage co-event background, whereas CD4/CD14 may be closer to helper T-cell/monocyte antigen-presentation context. Accordingly, CD3/CD14 is treated as a prespecified representative readout, not a uniquely specific pair."
    ),
    (
        "Clinically, this framework may help explain why whole-blood RNA signatures and monocyte HLA-DR protein or cytometry readouts can appear discordant within the same syndrome, without implying bedside prediction or treatment selection."
    ): (
        "Clinically, this interpretation may help explain why whole-blood RNA signatures and monocyte HLA-DR protein or cytometry readouts can appear discordant within the same syndrome, without implying bedside prediction or treatment selection."
    ),
    (
        "The validation strategy reflects the inherent constraints of public-data computational research. Validation was addressed through cross-cohort recurrence, continuous-spectrum analysis, composition sensitivity, paired RNA-CyTOF bridging, participant-level robustness, raw-FCS artifact-control analyses, donor-level single-cell context, pathway context, and explicit claim-boundary auditing. "
        "These analyses strengthen the evidence chain but do not replace the requirement for prospective biological validation. The resulting bounded working interpretation is summarized in Figure 6, with the evidence hierarchy, claim-boundary matrix and reproducibility traceability summarized in Supplementary Figures S1 and S2."
    ): (
        "Because this is a public-data computational study, robustness had to come from converging checks rather than a single validation experiment. "
        "We therefore combined cross-cohort recurrence, continuous-spectrum analysis, composition sensitivity, paired RNA-CyTOF bridging, participant-level robustness, raw-FCS artifact controls, donor-level single-cell context, pathway context and claim-boundary auditing. "
        "These checks strengthen the evidence chain but do not replace prospective biological validation. Figure 6 summarizes the working interpretation, and Supplementary Figures S1 and S2 summarize the evidence hierarchy, claim boundaries and reproducibility traceability."
    ),
    # AI statement
    (
        "ChatGPT (OpenAI; ChatGPT web/app) was used for language polishing and code correction. "
        "The authors reviewed and verified all analyses, figures, references, interpretations, and conclusions, and take full responsibility for the content. No AI tool is listed as an author."
    ): (
        "ChatGPT (OpenAI, ChatGPT web application, GPT-5.5 Thinking, accessed June 2026) was used for language polishing and code correction. "
        "The authors reviewed and verified all analyses, figures, references, interpretations, and conclusions, and take full responsibility for the content. No AI tool is listed as an author."
    ),
    # Figure legends
    (
        "Figure 1. Study architecture and claim-boundary map. (A) Evidence hierarchy from public bulk RNA recurrence, continuous case-only immune spectrum, representative residual co-event analysis, paired COMBAT RNA-CyTOF bridging, donor-level single-cell context and exploratory 28-day death-status anchoring. "
        "(B) Locked interpretation: inflammatory/MHC-II decoupling is associated with a representative residual co-event signal within a broader marker-pair co-event background in paired sepsis samples. "
        "(C) Claim boundaries separating allowed cross-cohort, paired, activation-state and computational interpretations from excluded causal mechanism, structural interaction, validated endotype, clinical prediction and therapeutic claims."
    ): (
        "Figure 1. Study architecture and claim-boundary map. (A) Evidence hierarchy from public bulk RNA recurrence, continuous case-only immune spectrum, representative residual co-event analysis, paired COMBAT RNA-CyTOF bridging, donor-level single-cell context and exploratory 28-day death-status anchoring. "
        "(B) Locked interpretation: inflammatory/MHC-II decoupling is associated with this residual co-event readout in paired sepsis samples. "
        "(C) Claim boundaries separating allowed cross-cohort, paired, activation-state and computational interpretations from excluded causal mechanism, structural interaction, validated endotype, clinical prediction and therapeutic claims."
    ),
    (
        "Figure 4. Raw-FCS marker-pair sensitivity and artifact-aware interpretation. (A) Prespecified marker-pair correlations across HLA-DR core, MHC-II/CD74 and RNA decoupling index. "
        "(B) All-pair raw-FCS marker-null distributions with CD3/CD14 positions. (C) Abundance-matched control-pair comparisons. (D) Lineage and QC residualized sensitivity. "
        "CD3/CD14 remains a representative residual co-event signal within a broader marker-pair co-event background; pair-level specificity remains unresolved, so the result is interpreted as an artifact-aware event-derived summary rather than a structural interaction."
    ): (
        "Figure 4. Raw-FCS marker-pair sensitivity. (A) Prespecified marker-pair correlations across HLA-DR core, MHC-II/CD74 and RNA decoupling index. "
        "(B) All-pair raw-FCS marker-null distributions with CD3/CD14 positions. (C) Abundance-matched control-pair comparisons. (D) Lineage and QC residualized sensitivity. "
        "CD3/CD14 remains a representative readout; pair-level specificity remains unresolved, so the result is interpreted as an event-derived summary rather than a structural interaction."
    ),
    (
        "Figure 5. COMBAT paired RNA-CyTOF bridge. (A) Pairing workflow from RNA-seq logCPM and strict event-QC CyTOF summaries to 129 matched participant-timepoint rows and 40 sepsis-only paired rows from 34 participants. "
        "(B-D) Sepsis-only paired associations linking a representative residual co-event abundance metric, operationalized by the prespecified CD3/CD14 pair, to MHC-II/CD74, HLA-DR core and RNA decoupling index. "
        "(E) Bootstrap confidence intervals. (F) RNA-CyTOF pairing-permutation null distributions. (G) Participant-aware sensitivity. The figure supports a paired public-data bridge while preserving association-only, non-causal and non-predictive claim boundaries."
    ): (
        "Figure 5. COMBAT paired RNA-CyTOF bridge. (A) Pairing workflow from RNA-seq logCPM and strict event-QC CyTOF summaries to 129 matched participant-timepoint rows and 40 sepsis-only paired rows from 34 participants. "
        "(B-D) Sepsis-only paired associations linking the prespecified residual co-event metric to MHC-II/CD74, HLA-DR core and RNA decoupling index. "
        "(E) Bootstrap confidence intervals. (F) RNA-CyTOF pairing-permutation null distributions. (G) Participant-aware sensitivity. The figure provides a paired public-data bridge under association-only, non-causal and non-predictive interpretation."
    ),
    (
        "Figure 6. Compartment-aware RNA-CyTOF immune-state decoupling map. (A) Whole-blood RNA layer summarizing inflammatory/S100 host-response increase, MHC-II/CD74/HLA-DR decrease and the decoupling index. "
        "(B) Paired COMBAT RNA-CyTOF bridge in sepsis-only samples. (C) CyTOF event-derived layer showing a representative residual co-event activation-state context with unresolved pair specificity. "
        "(D) Claim boundary: the diagram is an association-only, artifact-aware working interpretation requiring prospective validation, not a causal mechanism, clinical validation or structural interaction model."
    ): (
        "Figure 6. Compartment-aware RNA-CyTOF immune-state decoupling map. (A) Whole-blood RNA layer summarizing inflammatory/S100 host-response increase, MHC-II/CD74/HLA-DR decrease and the decoupling index. "
        "(B) Paired COMBAT RNA-CyTOF bridge in sepsis-only samples. (C) CyTOF event-derived layer showing residual co-event activation-state context with unresolved pair specificity. "
        "(D) Claim boundary: association-only working interpretation requiring prospective validation, not a causal mechanism, clinical validation or structural interaction model."
    ),
}


def run(cmd: list[str], *, check: bool = True, env: dict[str, str] | None = None) -> subprocess.CompletedProcess[str]:
    proc = subprocess.run(
        cmd,
        cwd=str(ROOT),
        env=env,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        check=False,
    )
    if check and proc.returncode != 0:
        raise RuntimeError(
            "Command failed:\n"
            + " ".join(cmd)
            + f"\nexit={proc.returncode}\nSTDOUT:\n{proc.stdout}\nSTDERR:\n{proc.stderr}"
        )
    return proc


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def rel(path: Path) -> str:
    try:
        return path.relative_to(ROOT).as_posix()
    except ValueError:
        return str(path)


def iter_paragraph_text(doc: Document) -> str:
    bits: list[str] = []
    for para in doc.paragraphs:
        bits.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    bits.append(para.text)
    return "\n".join(bits)


def patch_docx() -> dict[str, object]:
    doc = Document(MAIN_DOCX)
    applied: list[dict[str, object]] = []
    missing: list[str] = []
    for old, new in REPLACEMENTS.items():
        matches = [para for para in doc.paragraphs if para.text == old]
        if len(matches) != 1:
            missing.append(old[:120])
            continue
        para = matches[0]
        para.text = new
        applied.append(
            {
                "old_prefix": old[:90],
                "new_prefix": new[:90],
                "old_words": len(old.split()),
                "new_words": len(new.split()),
            }
        )
    if missing:
        raise RuntimeError(f"Expected replacement paragraphs not found: {missing}")
    doc.save(MAIN_DOCX)
    UPLOAD.mkdir(parents=True, exist_ok=True)
    shutil.copy2(MAIN_DOCX, UPLOAD_MAIN_DOCX)
    return {"applied_count": len(applied), "applied": applied, "docx_sha256": sha256(MAIN_DOCX)}


def count_docx_drawings(docx_path: Path) -> dict[str, object]:
    with zipfile.ZipFile(docx_path) as zf:
        names = zf.namelist()
        xml = "\n".join(
            zf.read(name).decode("utf-8", errors="ignore")
            for name in names
            if name.startswith("word/") and name.endswith(".xml")
        )
        media = [name for name in names if name.startswith("word/media/")]
    return {
        "drawing_tags": len(re.findall(r"<w:drawing[\s>]", xml)),
        "media_count": len(media),
        "tracked_change_tags": sum(
            len(re.findall(pattern, xml))
            for pattern in (r"<w:ins[\s>]", r"<w:del[\s>]", r"<w:moveFrom[\s>]", r"<w:moveTo[\s>]")
        ),
        "comment_parts": [name for name in names if name.startswith("word/comments")],
    }


def text_qa(docx_path: Path) -> dict[str, object]:
    text = iter_paragraph_text(Document(docx_path))
    phrases = {
        LONG_PHRASE: text.count(LONG_PHRASE),
        "representative residual co-event signal": text.count("representative residual co-event signal"),
        "broader marker-pair co-event background": text.count("broader marker-pair co-event background"),
        "artifact-aware": text.count("artifact-aware"),
        "bounded": text.count("bounded"),
        "claim-boundary": text.count("claim-boundary"),
        "claim boundaries": text.count("claim boundaries"),
        "structural interaction": text.count("structural interaction"),
    }
    forbidden = {
        "journal_submission_trace": "designed for submission to Frontiers",
        "journal_named_validation": "Frontiers-relevant validation",
        "reviewer_facing_artifact_label": "Artifact boundary for reviewers",
        "strong_coevent_remodeling": "CD3/CD14 co-event remodeling",
        "linked_immunoparalysis_phrase": "CD3/CD14-linked immunoparalysis",
        "specific_signal_phrase": "CD3/CD14-specific signal",
    }
    contribution = next((p.text for p in Document(docx_path).paragraphs if p.text.startswith("Sepsis often combines")), "")
    abstract = next((p.text for p in Document(docx_path).paragraphs if p.text.startswith("Background: Whole-blood")), "")
    ai_statement = next((p.text for p in Document(docx_path).paragraphs if p.text.startswith("ChatGPT (OpenAI")), "")
    return {
        "docx": rel(docx_path),
        "sha256": sha256(docx_path),
        "phrase_counts": phrases,
        "forbidden_hits": {label: text.count(term) for label, term in forbidden.items()},
        "square_numeric_citation_count": len(re.findall(r"\[\d+(?:-\d+)?(?:,\d+(?:-\d+)?)*\]", text)),
        "contribution_words": len(contribution.split()),
        "abstract_words": len(abstract.split()),
        "ai_statement": ai_statement,
        **count_docx_drawings(docx_path),
    }


def render_docx(docx_path: Path, variant: str) -> dict[str, object]:
    out_dir = QA / f"render_{variant}"
    if out_dir.exists():
        shutil.rmtree(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    env = os.environ.copy()
    env["TMPDIR"] = "/private/tmp"
    if SOFFICE_APP.exists():
        env["PATH"] = f"{SOFFICE_APP}:{env.get('PATH', '')}"
    proc = run([sys.executable, str(DOCX_RENDERER), str(docx_path), "--output_dir", str(out_dir), "--emit_pdf"], check=False, env=env)
    pdfs = sorted(out_dir.glob("*.pdf"))
    pngs = sorted(out_dir.glob("page-*.png"), key=lambda p: int(p.stem.split("-")[1]))
    pdf_out = UPLOAD / f"{docx_path.stem}.pdf"
    success = proc.returncode == 0 and bool(pdfs) and bool(pngs)
    if success:
        shutil.copy2(pdfs[0], pdf_out)
    return {
        "variant": variant,
        "success": success,
        "returncode": proc.returncode,
        "rendered_png_count": len(pngs),
        "pdf": rel(pdf_out) if pdf_out.exists() else None,
        "pdf_sha256": sha256(pdf_out) if pdf_out.exists() else None,
        "png_dir": rel(out_dir),
        "stdout_tail": proc.stdout[-1200:],
        "stderr_tail": proc.stderr[-1200:],
    }


def make_contact_sheet(render_dir: Path, out_path: Path) -> None:
    pages = sorted(render_dir.glob("page-*.png"), key=lambda p: int(p.stem.split("-")[1]))
    thumbs = []
    for page in pages:
        img = Image.open(page).convert("RGB")
        img.thumbnail((260, 200))
        tile = Image.new("RGB", (290, 238), "white")
        tile.paste(img, ((290 - img.width) // 2, 8))
        thumbs.append((tile, page.stem))
    cols = 3
    rows = max(1, (len(thumbs) + cols - 1) // cols)
    sheet = Image.new("RGB", (cols * 290, rows * 238), "white")
    draw = ImageDraw.Draw(sheet)
    for idx, (tile, label) in enumerate(thumbs):
        x = (idx % cols) * 290
        y = (idx // cols) * 238
        sheet.paste(tile, (x, y))
        draw.text((x + 8, y + 216), label, fill=(0, 0, 0))
    sheet.save(out_path, dpi=(150, 150))


def write_manifest_and_zip() -> dict[str, object]:
    files = []
    for path in sorted(p for p in OUT.rglob("*") if p.is_file()):
        rel_path = path.relative_to(OUT).as_posix()
        if rel_path.startswith("02_advisor_review/"):
            continue
        files.append((rel_path, path))

    manifest = QA / "FI_submission_revised_manifest_sha256.csv"
    with manifest.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=["relative_path", "size_bytes", "sha256"])
        writer.writeheader()
        for rel_path, path in files:
            writer.writerow({"relative_path": rel_path, "size_bytes": path.stat().st_size, "sha256": sha256(path)})

    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for rel_path, path in files:
            zf.write(path, f"{OUT.name}/{rel_path}")
    proc = run(["unzip", "-t", str(ZIP_PATH)], check=False)
    (QA / "zip_test.txt").write_text(proc.stdout + proc.stderr, encoding="utf-8")
    return {
        "manifest": rel(manifest),
        "file_count": len(files),
        "zip": rel(ZIP_PATH),
        "zip_success": proc.returncode == 0,
        "zip_sha256": sha256(ZIP_PATH),
        "zip_size_bytes": ZIP_PATH.stat().st_size,
    }


def update_readme() -> None:
    readme = OUT / "README.md"
    text = readme.read_text(encoding="utf-8") if readme.exists() else "# FI submission revised package\n"
    note = "- Applied final light language polish to reduce repeated defensive phrasing while preserving association-only and artifact-control boundaries."
    if note not in text:
        marker = "## Revision Notes\n"
        if marker in text:
            text = text.replace(marker, marker + note + "\n", 1)
        else:
            text += "\n## Revision Notes\n" + note + "\n"
        readme.write_text(text, encoding="utf-8")


def main() -> None:
    for path in [MAIN_DOCX, DOCX_RENDERER]:
        if not path.exists():
            raise FileNotFoundError(path)
    QA.mkdir(parents=True, exist_ok=True)
    patch_result = patch_docx()
    render_result = render_docx(MAIN_DOCX, "v21_main_text_only_revised")
    if not render_result["success"]:
        raise RuntimeError("Main DOCX render failed")
    make_contact_sheet(ROOT / render_result["png_dir"], QA / "v21_main_text_only_revised_contact_sheet.png")
    update_readme()
    zip_result = write_manifest_and_zip()

    report_path = QA / "final_light_language_polish_QA_report.json"
    report = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "scope": [
            "Contribution to the Field",
            "Abstract",
            "final Introduction paragraph",
            "first six Discussion paragraphs",
            "Figure 1, 4, 5 and 6 legends",
            "Generative AI statement",
        ],
        "patch_result": patch_result,
        "text_QA": text_qa(MAIN_DOCX),
        "render_QA": render_result,
        "zip_result": zip_result,
    }
    report_path.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")

    revision_report = QA / "revision_QA_report.json"
    if revision_report.exists():
        data = json.loads(revision_report.read_text(encoding="utf-8"))
        data["final_light_language_polish"] = {
            "report": rel(report_path),
            "main_docx_sha256": sha256(MAIN_DOCX),
            "main_pdf_sha256": render_result["pdf_sha256"],
            "phrase_counts": report["text_QA"]["phrase_counts"],
            "ai_statement_updated": "GPT-5.5 Thinking" in report["text_QA"]["ai_statement"],
        }
        data["zip_result"] = {
            "zip": rel(ZIP_PATH),
            "success": zip_result["zip_success"],
            "sha256": zip_result["zip_sha256"],
            "size_bytes": zip_result["zip_size_bytes"],
        }
        revision_report.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")

    print(
        json.dumps(
            {
                "docx": str(MAIN_DOCX),
                "upload_docx": str(UPLOAD_MAIN_DOCX),
                "pdf": str(UPLOAD / f"{MAIN_DOCX.stem}.pdf"),
                "qa_report": str(report_path),
                "rendered_pages": render_result["rendered_png_count"],
                "phrase_counts": report["text_QA"]["phrase_counts"],
                "zip_success": zip_result["zip_success"],
            },
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
