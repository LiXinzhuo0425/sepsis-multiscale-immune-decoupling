#!/usr/bin/env python3
"""Align FI manuscript legends with final figures and add Supplementary Figures S1-S2.

This is a packaging/formatting script only. It does not rerun analyses and does not
change the six user-approved main figures.
"""

from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET
from xml.sax.saxutils import escape

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.shared import Pt
from PIL import Image, ImageChops, ImageDraw


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "06_FI_SUBMISSION_READY_20260620"
UPLOAD = OUT / "01_upload_ready"
VECTOR = OUT / "02_vector_backup"
SOURCE = OUT / "03_source_data_and_attribution"
QA = OUT / "04_QA_and_manifests"
ZIP_PATH = ROOT / "FI_submission_ready_20260620.zip"

ORIGINAL_DOCX = ROOT / "00_FINAL_WORD/Frontiers_Immunology_V17_FINAL_TEXT_ONLY.docx"
REVISED_DOCX = ROOT / "00_FINAL_WORD/Frontiers_Immunology_V18_FI_figure_aligned.docx"
DOCX_RENDERER = Path(
    "<CODEX_DOCUMENTS_RENDER_DOCX>"
)
SOFFICE_APP = Path("/Applications/LibreOffice.app/Contents/MacOS")

DPI = 600
SUPP_W_PT = 510.2
SUPP_H_PT = 425.2

TRACE = ROOT / "03_TRACEABLE_RESULTS"
STAGE7 = TRACE / "stage_outputs_tables/03_REPRODUCIBLE_RESULTS/stage7_review_hardening"
STAGE7_QA = TRACE / "QA_and_audit_reports/03_REPRODUCIBLE_RESULTS/stage7_review_hardening"
FIG_SOURCE = TRACE / "figure_source_data_tables/03_REPRODUCIBLE_RESULTS"
V17_FIG_SOURCE = FIG_SOURCE / "frontiers_v17_visual_engineering_polish/figure_source_data"
STAGE7_FIG_SOURCE = FIG_SOURCE / "stage7_review_hardening"


@dataclass(frozen=True)
class SuppFig:
    stem: str
    title: str
    caption: str
    source_files: tuple[Path, ...]


S1 = SuppFig(
    stem="Supplementary_Figure_S1",
    title="Evidence hierarchy and claim-boundary matrix",
    caption=(
        "Supplementary Figure S1. Evidence hierarchy and claim-boundary matrix. "
        "The schematic summarizes the ordered evidence layers and separates allowed "
        "association-only interpretations from excluded causal, structural, endotype, "
        "clinical-prediction and therapeutic claims."
    ),
    source_files=(
        STAGE7 / "stage7_claim_boundary_matrix.csv",
        STAGE7_QA / "stage7_dataset_audit_master_table.csv",
        V17_FIG_SOURCE / "Figure1_design_claim_boundary_source.csv",
        STAGE7 / "stage7_clinical_anchor_map.csv",
    ),
)

S2 = SuppFig(
    stem="Supplementary_Figure_S2",
    title="Data-source and reproducibility traceability map",
    caption=(
        "Supplementary Figure S2. Data-source and reproducibility traceability map. "
        "The schematic links public data layers, source tables, formal supplementary "
        "tables, figure exports, QA checks, checksums and the final upload-ready package "
        "without exposing private local paths."
    ),
    source_files=(
        STAGE7_FIG_SOURCE / "stage7_figure_rebuild_plan.csv",
        STAGE7 / "stage7_supplementary_table_index.csv",
        QA / "figure_QA_summary.csv",
        QA / "frontiers_upload_figure_manifest_sha256.csv",
        QA / "FI_submission_ready_manifest_sha256.csv",
    ),
)

SUPP_FIGS = (S1, S2)


def run(cmd: list[str], *, cwd: Path = ROOT, check: bool = True, env: dict[str, str] | None = None) -> subprocess.CompletedProcess[str]:
    proc = subprocess.run(
        cmd,
        cwd=str(cwd),
        env=env,
        check=False,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    if check and proc.returncode != 0:
        raise RuntimeError(
            "Command failed:\n"
            + " ".join(cmd)
            + f"\nexit={proc.returncode}\nSTDOUT:\n{proc.stdout}\nSTDERR:\n{proc.stderr}"
        )
    return proc


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def rel(path: Path, base: Path = ROOT) -> str:
    try:
        return path.relative_to(base).as_posix()
    except ValueError:
        return str(path)


def clear_set(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def compact_legend_layout(doc: Document) -> None:
    in_figure_legends = False
    in_supp_legends = False
    for para in list(doc.paragraphs):
        text = para.text.strip()
        if text == "Figure Legends":
            in_figure_legends = True
            pf = para.paragraph_format
            pf.keep_with_next = False
            pf.keep_together = False
            pf.page_break_before = False
            pf.space_before = Pt(12)
            pf.space_after = Pt(10)
            continue
        if text == "Tables":
            in_figure_legends = False
        if text == "Supplementary Figure Legends":
            in_supp_legends = True
            pf = para.paragraph_format
            pf.keep_with_next = False
            pf.keep_together = False
            pf.page_break_before = False
            pf.space_before = Pt(12)
            pf.space_after = Pt(10)
            continue
        if text == "Supplementary Materials":
            in_supp_legends = False

        if in_figure_legends or in_supp_legends:
            if not text:
                delete_paragraph(para)
                continue
            pf = para.paragraph_format
            pf.keep_with_next = False
            pf.keep_together = False
            pf.page_break_before = False
            pf.space_before = Pt(0)
            pf.space_after = Pt(6)
            pf.line_spacing = 1.0


def iter_doc_text(doc: Document) -> Iterable[str]:
    for p in doc.paragraphs:
        yield p.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p.text


def revise_docx() -> dict[str, object]:
    doc = Document(ORIGINAL_DOCX)

    replacements = {
        "Here we investigated whether public RNA and cytometry data could support a bounded, compartment-aware model of sepsis-associated immunoparalysis. The study design and claim boundaries are summarized in Figure 1.": (
            "Here we investigated whether public RNA and cytometry data could support a bounded, compartment-aware model of sepsis-associated immunoparalysis. The study architecture and claim boundaries are summarized in Figure 1."
        ),
        "Our evidence hierarchy proceeds from cross-cohort bulk RNA recurrence to continuous-spectrum case-only computational grouping, residual CD3/CD14 co-event signal remodeling and pair-specificity hardening, a paired COMBAT RNA-CyTOF bridge, donor-level single-cell contextual support, and an exploratory 28-day mortality clinical anchor.": (
            "Our evidence hierarchy proceeds from cross-cohort bulk RNA recurrence to a continuous case-only immune-decoupling spectrum, raw-FCS pair-specificity hardening of the residual CD3/CD14 co-event signal, a paired COMBAT RNA-CyTOF bridge, donor-level single-cell contextual support, pathway context and an exploratory 28-day mortality clinical anchor."
        ),
        "In the public COMBAT CELLxGENE PBMC dataset, donor-level monocyte/MNP pseudobulk analysis compared 23 sepsis donors with 10 healthy volunteer controls and showed lower monocyte MHC-II/HLA-DR context in sepsis (Supplementary Figure S1), higher S100 and inflammatory context, and inverse sample-level correlation between monocyte MHC-II/HLA-DR context and S100/inflammatory context.": (
            "In the public COMBAT CELLxGENE PBMC dataset, donor-level monocyte/MNP pseudobulk analysis compared 23 sepsis donors with 10 healthy volunteer controls and showed lower monocyte MHC-II/HLA-DR context in sepsis, higher S100 and inflammatory context, and inverse sample-level correlation between monocyte MHC-II/HLA-DR context and S100/inflammatory context (Supplementary Table S18)."
        ),
        "These analyses provide cellular and pathway plausibility, not paired validation or cell-cell interaction validation (Supplementary Figure S1 and Supplementary Table S18).": (
            "These analyses provide cellular and pathway plausibility, not paired validation or cell-cell interaction validation (Supplementary Table S18)."
        ),
        "The resulting bounded working interpretation is summarized in Figure 6.": (
            "The resulting bounded working interpretation is summarized in Figure 6, with the evidence hierarchy, claim-boundary matrix and reproducibility traceability summarized in Supplementary Figures S1 and S2."
        ),
        "The public single-cell analysis provides donor-level monocyte context (Supplementary Figure S1) but not paired RNA-CyTOF validation.": (
            "The public single-cell analysis provides donor-level monocyte context (Supplementary Table S18) but not paired RNA-CyTOF validation."
        ),
        "Formal supplementary tables are provided as Supplementary Tables S1 to S22.": (
            "Formal supplementary materials include Supplementary Figures S1-S2 and Supplementary Tables S1 to S22."
        ),
        "This appendix is appended after the manuscript text. It is text/table only: no figures or embedded image media are added.": (
            "This appendix is appended after the manuscript text. It is text/table only; Supplementary Figures S1-S2 are provided as separate image files and are not embedded in this Word document."
        ),
    }

    figure_legends = {
        "Figure 1.": (
            "Figure 1. Study architecture and claim-boundary map. (A) Evidence hierarchy from public bulk RNA recurrence, continuous case-only immune spectrum, residual co-event remodeling, paired COMBAT RNA-CyTOF bridging, donor-level single-cell context and exploratory Death28 anchoring. (B) Locked interpretation: inflammatory/MHC-II decoupling is associated with a prespecified residual CD3/CD14 co-event signal in paired sepsis samples. (C) Claim boundaries separating allowed cross-cohort, paired, activation-state and computational interpretations from excluded causal mechanism, physical CD3/CD14 complex, validated endotype, clinical prediction and therapeutic claims."
        ),
        "Figure 2.": (
            "Figure 2. Cross-cohort inflammatory/MHC-II decoupling in public sepsis RNA cohorts. (A) Signature direction atlas across control-available cohorts. (B) Primary inverse coupling of the six-gene host-response panel with MHC-II/CD74 and HLA-DR context across cohorts. (C) Leave-one-cohort sensitivity showing preserved direction after sequential cohort removal. (D) Gene-count-matched random-signature null intervals; these nulls are not expression-level matched and are used as robustness context rather than validation of a classifier."
        ),
        "Figure 3.": (
            "Figure 3. Continuous-spectrum case-only immune decoupling and pathway context. (A) Case-only immune-decoupling spectrum in 499 sepsis cases, with k=3 shown as a descriptive discretization of a continuous PC1 gradient. (B) PC1 alignment with host-response, myeloid, immunometabolic, MHC-II/CD74, HLA-DR, interferon/antigen-presentation and adaptive/T-cell programs. (C) Cross-cohort centroid-transfer stability trained from the largest cohort. (D) Ranked pathway context from recoverable public cohorts. These results do not establish validated clinical endotypes or causal pathway effects."
        ),
        "Figure 4.": (
            "Figure 4. Raw-FCS pair-specificity hardening and artifact boundary. (A) Prespecified marker-pair correlations across HLA-DR core, MHC-II/CD74 and RNA decoupling index. (B) All-pair raw-FCS marker-null distributions with CD3/CD14 positions. (C) Abundance-matched control-pair comparisons. (D) Lineage and QC residualized sensitivity. CD3/CD14 remains a prespecified representative residual co-event signal; pair-level specificity remains unresolved, so the result is interpreted as an artifact-aware event-derived summary rather than a structural interaction."
        ),
        "Figure 5.": (
            "Figure 5. COMBAT paired RNA-CyTOF bridge. (A) Pairing workflow from RNA-seq logCPM and strict event-QC CyTOF summaries to 129 matched participant-timepoint rows and 40 sepsis-only paired rows from 34 participants. (B-D) Sepsis-only paired associations linking residual CD3/CD14 co-event abundance to MHC-II/CD74, HLA-DR core and RNA decoupling index. (E) Bootstrap confidence intervals. (F) RNA-CyTOF pairing-permutation null distributions. (G) Participant-aware sensitivity. The figure supports a paired public-data bridge while preserving association-only, non-causal and non-predictive claim boundaries."
        ),
        "Figure 6.": (
            "Figure 6. Compartment-aware RNA-CyTOF immune-state decoupling map. (A) Whole-blood RNA layer summarizing inflammatory/S100 host-response increase, MHC-II/CD74/HLA-DR decrease and the decoupling index. (B) Paired COMBAT RNA-CyTOF bridge in sepsis-only samples. (C) CyTOF event-derived layer showing residual CD3/CD14 co-event activation-state context and unresolved pair specificity. (D) Claim boundary: the diagram is an association-only, artifact-aware working interpretation requiring prospective validation, not a causal mechanism, clinical validation or structural interaction model."
        ),
        "Supplementary Figure S1.": S1.caption,
    }

    replacement_hits = {key: 0 for key in replacements}
    legend_hits = {key: 0 for key in figure_legends}
    s1_para: Paragraph | None = None

    for para in doc.paragraphs:
        text = para.text
        for old, new in replacements.items():
            if old in text:
                text = text.replace(old, new)
                replacement_hits[old] += 1
        for prefix, new in figure_legends.items():
            if text.startswith(prefix):
                text = new
                legend_hits[prefix] += 1
                if prefix == "Supplementary Figure S1.":
                    s1_para = para
                break
        if text != para.text:
            clear_set(para, text)

    for para in doc.paragraphs:
        if para.text.strip() == "Supplementary Figure S1":
            clear_set(para, "Supplementary Figure Legends")

    if s1_para is None:
        raise RuntimeError("Could not locate Supplementary Figure S1 legend paragraph")
    paragraph_after(s1_para, "")
    paragraph_after(s1_para, S2.caption)
    compact_legend_layout(doc)

    doc.save(REVISED_DOCX)
    stale_upload_docx = UPLOAD / ORIGINAL_DOCX.name
    if stale_upload_docx.exists():
        stale_upload_docx.unlink()
    shutil.copy2(REVISED_DOCX, UPLOAD / REVISED_DOCX.name)

    return {
        "revised_docx": rel(REVISED_DOCX),
        "upload_docx": rel(UPLOAD / REVISED_DOCX.name),
        "replacement_hits": replacement_hits,
        "legend_hits": legend_hits,
    }


def svg_text(x: float, y: float, text: str, size: float = 8.0, weight: int | str = 400, fill: str = "#1f2933", anchor: str = "start") -> str:
    return (
        f'<text x="{x:.1f}" y="{y:.1f}" font-family="Arial, Helvetica, sans-serif" '
        f'font-size="{size:.1f}" font-weight="{weight}" fill="{fill}" text-anchor="{anchor}">'
        f"{escape(text)}</text>"
    )


def wrapped_text(x: float, y: float, text: str, width_chars: int, size: float = 7.2, line_gap: float = 9.0, weight: int | str = 400, fill: str = "#1f2933") -> tuple[str, float]:
    words = text.split()
    lines: list[str] = []
    cur = ""
    for word in words:
        candidate = word if not cur else f"{cur} {word}"
        if len(candidate) <= width_chars:
            cur = candidate
        else:
            if cur:
                lines.append(cur)
            cur = word
    if cur:
        lines.append(cur)
    parts = []
    yy = y
    for line in lines:
        parts.append(svg_text(x, yy, line, size=size, weight=weight, fill=fill))
        yy += line_gap
    return "\n".join(parts), yy


def rect(x: float, y: float, w: float, h: float, fill: str, stroke: str = "#c8d0d8", rx: float = 6.0, sw: float = 1.0) -> str:
    return f'<rect x="{x:.1f}" y="{y:.1f}" width="{w:.1f}" height="{h:.1f}" rx="{rx:.1f}" fill="{fill}" stroke="{stroke}" stroke-width="{sw:.1f}"/>'


def line(x1: float, y1: float, x2: float, y2: float, stroke: str = "#8aa0b5", sw: float = 1.4, marker: bool = False) -> str:
    m = ' marker-end="url(#arrow)"' if marker else ""
    return f'<line x1="{x1:.1f}" y1="{y1:.1f}" x2="{x2:.1f}" y2="{y2:.1f}" stroke="{stroke}" stroke-width="{sw:.1f}"{m}/>'


def bullet(x: float, y: float, text: str, width_chars: int = 34, size: float = 6.9, fill: str = "#1f2933") -> tuple[str, float]:
    dot = f'<circle cx="{x:.1f}" cy="{y - 2.4:.1f}" r="2.0" fill="#497a8f"/>'
    body, yy = wrapped_text(x + 7, y, text, width_chars, size=size, line_gap=size + 1.9, fill=fill)
    return dot + "\n" + body, yy + 1.5


def svg_shell(title: str, body: str, desc: str = "") -> str:
    return f'''<?xml version="1.0" encoding="UTF-8"?>
<svg xmlns="http://www.w3.org/2000/svg" width="{SUPP_W_PT}pt" height="{SUPP_H_PT}pt" viewBox="0 0 {SUPP_W_PT} {SUPP_H_PT}">
<desc>{escape(desc or title)}</desc>
<defs>
  <marker id="arrow" markerWidth="8" markerHeight="8" refX="7" refY="3.5" orient="auto">
    <path d="M0,0 L7,3.5 L0,7 Z" fill="#8aa0b5"/>
  </marker>
  <style>
    text {{ dominant-baseline: alphabetic; }}
  </style>
</defs>
<rect x="0" y="0" width="{SUPP_W_PT}" height="{SUPP_H_PT}" fill="#ffffff"/>
{svg_text(18, 25, title, size=13.2, weight=700, fill="#17212b")}
{body}
</svg>
'''


def build_s1_svg() -> str:
    parts: list[str] = []
    parts.append(svg_text(18, 43, "Evidence layers are ordered by proximity to the manuscript claim; all claims remain association-only.", 7.8, 400, "#4b5966"))
    parts.append(rect(18, 58, 230, 304, "#f7fafc", "#d4dde6", 8))
    parts.append(svg_text(30, 78, "A  Evidence hierarchy", 10, 700, "#102a43"))
    layers = [
        ("1", "Public bulk RNA recurrence", "7 cohorts; 660 samples, 499 sepsis; recurrent inflammatory/MHC-II inverse coupling"),
        ("2", "Case-only immune spectrum", "PC1 gradient 55.9%; k=3 descriptive; 221/499 high-decoupling cases"),
        ("3", "Raw-FCS hardening", "Residual CD3/CD14 signal; negative-pair, all-pair and QC sensitivity; specificity unresolved"),
        ("4", "Paired COMBAT bridge", "129 matched rows; 40 sepsis-only rows from 34 participants; paired association layer"),
        ("5", "Context layers", "Donor-level single-cell and pathway context; plausibility support only"),
        ("6", "Death28 anchor", "40 rows; 11 events; exploratory clinical anchor only"),
    ]
    y = 93
    for idx, heading, desc in layers:
        parts.append(rect(30, y, 206, 36, "#ffffff", "#d9e2ec", 5))
        parts.append(f'<circle cx="45" cy="{y + 18:.1f}" r="10.2" fill="#dceef7" stroke="#8bb7c7" stroke-width="1"/>')
        parts.append(svg_text(45, y + 21, idx, 8.0, 700, "#123b52", "middle"))
        parts.append(svg_text(61, y + 14, heading, 7.8, 700, "#17212b"))
        txt, _ = wrapped_text(61, y + 25, desc, 46, 6.2, 7.4, fill="#52616f")
        parts.append(txt)
        if idx != "6":
            parts.append(line(45, y + 36, 45, y + 43, "#9fb3c8", 1.0, True))
        y += 43

    parts.append(rect(263, 58, 229, 304, "#f7fafc", "#d4dde6", 8))
    parts.append(svg_text(275, 78, "B  Claim-boundary matrix", 10, 700, "#102a43"))
    parts.append(rect(276, 92, 94, 25, "#e3f4ec", "#96d4b7", 5))
    parts.append(rect(384, 92, 94, 25, "#fff0e8", "#efb08b", 5))
    parts.append(svg_text(323, 108, "Allowed", 8.5, 700, "#1f6f4a", "middle"))
    parts.append(svg_text(431, 108, "Excluded", 8.5, 700, "#9a4d24", "middle"))

    allowed = [
        "Cross-cohort recurrent association",
        "Matched public RNA-CyTOF bridge",
        "Activation-state event context",
        "Computational spectrum",
        "Exploratory clinical anchor",
    ]
    excluded = [
        "Causal mechanism",
        "Physical CD3/CD14 complex",
        "Validated clinical endotype",
        "Clinical prediction model",
        "Therapeutic recommendation",
    ]
    y0 = 133
    for i, text in enumerate(allowed):
        row_y = y0 + i * 34
        parts.append(rect(276, row_y - 15, 94, 25, "#ffffff", "#d9e2ec", 4))
        parts.append(rect(384, row_y - 15, 94, 25, "#ffffff", "#d9e2ec", 4))
        t1, _ = wrapped_text(283, row_y - 2, text, 21, 6.4, 7.2, fill="#264653")
        t2, _ = wrapped_text(391, row_y - 2, excluded[i], 21, 6.4, 7.2, fill="#6b3f2a")
        parts.append(t1)
        parts.append(t2)

    parts.append(rect(276, 314, 202, 32, "#eef5ff", "#b6c7df", 5))
    note, _ = wrapped_text(
        286,
        329,
        "Use S1 as an audit map: it explains how the figures support a bounded interpretation but does not add a new analytical result.",
        58,
        6.5,
        7.4,
        fill="#40566b",
    )
    parts.append(note)
    parts.append(svg_text(18, 389, "Supplementary Figure S1. Evidence hierarchy and claim-boundary matrix.", 6.6, 400, "#4b5966"))
    parts.append(svg_text(18, 400, "Schematic separates allowed association-only interpretations from excluded causal, clinical and therapeutic claims.", 6.6, 400, "#4b5966"))
    return svg_shell("Supplementary Figure S1. Evidence hierarchy and claim-boundary matrix", "\n".join(parts), S1.caption)


def build_s2_svg() -> str:
    parts: list[str] = []
    parts.append(svg_text(18, 43, "Traceability is shown as public/source layer to figure/table outputs to export QA and checksums.", 7.8, 400, "#4b5966"))

    cols = [
        (18, 64, 106, "A  Public/source layers", "#eef7f6"),
        (144, 64, 106, "B  Source tables", "#f7fafc"),
        (270, 64, 106, "C  Manuscript outputs", "#fffaf0"),
        (396, 64, 96, "D  QA package", "#f5f1ff"),
    ]
    for x, y, w, title, fill in cols:
        parts.append(rect(x, y, w, 290, fill, "#d4dde6", 8))
        parts.append(svg_text(x + 8, y + 19, title, 8.3, 700, "#102a43"))

    source_items = [
        "7 public bulk RNA cohorts",
        "COMBAT paired RNA-CyTOF",
        "Raw-FCS sensitivity layer",
        "CELLxGENE/pathway context",
        "Death28 exploratory anchor",
    ]
    y = 104
    for item in source_items:
        parts.append(rect(28, y - 15, 84, 24, "#ffffff", "#d9e2ec", 4))
        txt, _ = wrapped_text(36, y - 2, item, 21, 6.3, 7.2, fill="#334e68")
        parts.append(txt)
        y += 39

    table_items = [
        "Dataset inventory and inclusion audit",
        "Figure source-data index",
        "Cytometry and raw-FCS controls",
        "COMBAT robustness tables",
        "Claim-boundary ledger",
    ]
    y = 104
    for item in table_items:
        parts.append(rect(154, y - 15, 84, 24, "#ffffff", "#d9e2ec", 4))
        txt, _ = wrapped_text(162, y - 2, item, 22, 6.3, 7.2, fill="#334e68")
        parts.append(txt)
        y += 39

    output_items = [
        "Main Figures 1-6",
        "Supplementary Figures S1-S2",
        "Supplementary Tables S1-S22",
        "Revised clean DOCX",
        "Vector backups",
    ]
    y = 104
    for item in output_items:
        parts.append(rect(280, y - 15, 84, 24, "#ffffff", "#e5d2a4", 4))
        txt, _ = wrapped_text(288, y - 2, item, 22, 6.3, 7.2, fill="#594a2e")
        parts.append(txt)
        y += 39

    qa_items = [
        "RGB TIFF, 600 dpi",
        "PDF/EPS/SVG backup",
        "No embedded raster in S1/S2",
        "SHA256 manifests",
        "Final zip test",
    ]
    y = 104
    for item in qa_items:
        parts.append(rect(405, y - 15, 76, 24, "#ffffff", "#cabffd", 4))
        txt, _ = wrapped_text(413, y - 2, item, 20, 6.3, 7.2, fill="#4c3d73")
        parts.append(txt)
        y += 39

    for yy in [104, 143, 182, 221, 260]:
        parts.append(line(116, yy - 3, 143, yy - 3, "#8aa0b5", 1.2, True))
        parts.append(line(242, yy - 3, 269, yy - 3, "#8aa0b5", 1.2, True))
        parts.append(line(368, yy - 3, 395, yy - 3, "#8aa0b5", 1.2, True))

    parts.append(rect(28, 313, 453, 27, "#ffffff", "#d4dde6", 5))
    note, _ = wrapped_text(
        38,
        329,
        "All displayed labels use public accession-level or package-level descriptors; private absolute paths are intentionally omitted from the graphic.",
        92,
        6.4,
        7.3,
        fill="#40566b",
    )
    parts.append(note)
    parts.append(svg_text(18, 389, "Supplementary Figure S2. Data-source and reproducibility traceability map.", 6.6, 400, "#4b5966"))
    parts.append(svg_text(18, 400, "Schematic links public/source layers to figure outputs, QA checks, checksums and the upload-ready package.", 6.6, 400, "#4b5966"))
    return svg_shell("Supplementary Figure S2. Data-source and reproducibility traceability map", "\n".join(parts), S2.caption)


def svg_stats(path: Path) -> dict[str, object]:
    root = ET.parse(path).getroot()
    local = lambda tag: tag.split("}", 1)[-1]
    viewbox = root.attrib.get("viewBox") or root.attrib.get("viewbox")
    if not viewbox:
        raise ValueError(f"Missing viewBox: {path}")
    return {
        "viewBox": viewbox,
        "text_elements": sum(1 for e in root.iter() if local(e.tag) == "text"),
        "image_elements": sum(1 for e in root.iter() if local(e.tag) == "image"),
        "bytes": path.stat().st_size,
        "sha256": sha256(path),
    }


def pdf_page_size(path: Path) -> str:
    proc = run(["pdfinfo", str(path)], check=False)
    if proc.returncode != 0:
        return f"ERROR: {proc.stderr.strip()}"
    for line_text in proc.stdout.splitlines():
        if line_text.startswith("Page size:"):
            return line_text.split(":", 1)[1].strip()
    return "UNKNOWN"


def pdf_font_rows(path: Path) -> int | str:
    proc = run(["pdffonts", str(path)], check=False)
    if proc.returncode != 0:
        return f"ERROR: {proc.stderr.strip()}"
    rows = [line_text for line_text in proc.stdout.splitlines() if line_text.strip()]
    return max(0, len(rows) - 2)


def tiff_info(path: Path) -> dict[str, object]:
    with Image.open(path) as img:
        dpi = img.info.get("dpi", (None, None))
        return {
            "mode": img.mode,
            "width_px": img.width,
            "height_px": img.height,
            "dpi_x": round(float(dpi[0]), 3) if dpi[0] else None,
            "dpi_y": round(float(dpi[1]), 3) if dpi[1] else None,
        }


def nonblank(path: Path) -> bool:
    with Image.open(path) as img:
        rgb = img.convert("RGB")
        bg = Image.new("RGB", rgb.size, "white")
        return ImageChops.difference(rgb, bg).getbbox() is not None


def export_supplementary_figures() -> tuple[list[dict[str, object]], list[dict[str, object]]]:
    vector_rows: list[dict[str, object]] = []
    qa_rows: list[dict[str, object]] = []
    builds = {
        S1.stem: build_s1_svg,
        S2.stem: build_s2_svg,
    }

    for spec in SUPP_FIGS:
        src_svg = VECTOR / f"{spec.stem}.source.svg"
        outlined_svg = VECTOR / f"{spec.stem}.outlined.svg"
        pdf = VECTOR / f"{spec.stem}.pdf"
        eps = VECTOR / f"{spec.stem}.eps"
        png = VECTOR / f"{spec.stem}_600dpi.png"
        tiff = VECTOR / f"{spec.stem}_600dpi.tiff"
        preview = VECTOR / f"{spec.stem}_preview.png"
        upload_tiff = UPLOAD / f"{spec.stem}.tiff"
        upload_preview = UPLOAD / f"{spec.stem}_preview.png"

        src_svg.write_text(builds[spec.stem](), encoding="utf-8")
        ET.parse(src_svg)

        run(["inkscape", str(src_svg), "--export-type=svg", "--export-text-to-path", f"--export-filename={outlined_svg}"])
        run(["inkscape", str(outlined_svg), "--export-type=pdf", f"--export-filename={pdf}"])
        run(["inkscape", str(outlined_svg), "--export-type=eps", f"--export-filename={eps}"])
        run([
            "inkscape",
            str(outlined_svg),
            "--export-type=png",
            f"--export-dpi={DPI}",
            "--export-background=white",
            "--export-background-opacity=1",
            f"--export-filename={png}",
        ])
        run([
            "inkscape",
            str(outlined_svg),
            "--export-type=png",
            "--export-width=1600",
            "--export-background=white",
            "--export-background-opacity=1",
            f"--export-filename={preview}",
        ])

        with Image.open(png) as img:
            rgb = img.convert("RGB")
            rgb.save(png, dpi=(DPI, DPI))
            rgb.save(tiff, dpi=(DPI, DPI), compression="tiff_lzw")
        shutil.copy2(tiff, upload_tiff)
        shutil.copy2(preview, upload_preview)

        stats = svg_stats(src_svg)
        out_stats = svg_stats(outlined_svg)
        tinfo = tiff_info(tiff)
        qa = {
            "supplementary_figure": spec.stem,
            "title": spec.title,
            **stats,
            "outlined_svg_text_elements": out_stats["text_elements"],
            "outlined_text_to_path_pass": out_stats["text_elements"] == 0,
            "pdf_page_size": pdf_page_size(pdf),
            "pdf_font_rows": pdf_font_rows(pdf),
            "tiff_mode": tinfo["mode"],
            "tiff_width_px": tinfo["width_px"],
            "tiff_height_px": tinfo["height_px"],
            "tiff_dpi_x": tinfo["dpi_x"],
            "tiff_dpi_y": tinfo["dpi_y"],
            "frontiers_min_dpi_pass": bool(tinfo["mode"] == "RGB" and (tinfo["dpi_x"] or 0) >= 300 and (tinfo["dpi_y"] or 0) >= 300),
            "preview_nonblank": nonblank(preview),
            "embedded_raster_count": stats["image_elements"],
            "caption_embedded_in_svg": spec.caption in src_svg.read_text(encoding="utf-8"),
        }
        qa_rows.append(qa)

        for role, path in {
            "source_svg": src_svg,
            "outlined_svg": outlined_svg,
            "pdf": pdf,
            "eps": eps,
            "png_600dpi": png,
            "tiff_600dpi": tiff,
            "preview_png": preview,
            "upload_tiff": upload_tiff,
            "upload_preview": upload_preview,
        }.items():
            vector_rows.append(
                {
                    "asset": spec.stem,
                    "role": role,
                    "path": rel(path),
                    "size_bytes": path.stat().st_size,
                    "sha256": sha256(path),
                }
            )

    return vector_rows, qa_rows


def copy_supplementary_source_data() -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    for spec in SUPP_FIGS:
        dest = SOURCE / spec.stem
        dest.mkdir(parents=True, exist_ok=True)
        mapping_rows = []
        for path in spec.source_files:
            if not path.exists():
                mapping_rows.append(
                    {
                        "source_file_label": path.name,
                        "copied_to": "",
                        "sha256": "",
                        "role": "expected source file not found",
                    }
                )
                continue
            dst = dest / path.name
            shutil.copy2(path, dst)
            mapping_rows.append(
                {
                    "source_file_label": path.name,
                    "copied_to": dst.relative_to(OUT).as_posix(),
                    "sha256": sha256(dst),
                    "role": "source/audit table used to build supplementary schematic",
                }
            )
            rows.append(
                {
                    "figure": spec.stem,
                    "source": path.name,
                    "copied_to": rel(dst),
                    "sha256": sha256(dst),
                }
            )
        mapping_rows.append(
            {
                "source_file_label": "external vector assets",
                "copied_to": "",
                "sha256": "",
                "role": "none used; all shapes are self-drawn SVG primitives",
            }
        )
        with (dest / "source_mapping_and_attribution.csv").open("w", newline="", encoding="utf-8") as handle:
            writer = csv.DictWriter(handle, fieldnames=["source_file_label", "copied_to", "sha256", "role"])
            writer.writeheader()
            writer.writerows(mapping_rows)
    return rows


def write_csv(path: Path, rows: list[dict[str, object]], fieldnames: list[str] | None = None) -> None:
    if fieldnames is None:
        keys: list[str] = []
        for row in rows:
            for key in row:
                if key not in keys:
                    keys.append(key)
        fieldnames = keys
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def update_source_manifest(supp_source_rows: list[dict[str, object]]) -> None:
    manifest = SOURCE / "source_data_copy_manifest.csv"
    existing: list[dict[str, object]] = []
    if manifest.exists():
        with manifest.open(newline="", encoding="utf-8") as handle:
            existing = list(csv.DictReader(handle))
    existing = [row for row in existing if not row.get("figure", "").startswith("Supplementary_Figure_")]
    existing.extend(supp_source_rows)
    existing = [
        {
            "figure": row.get("figure", ""),
            "source": row.get("source", ""),
            "copied_to": row.get("copied_to", ""),
            "sha256": row.get("sha256", ""),
        }
        for row in existing
    ]
    write_csv(manifest, existing, ["figure", "source", "copied_to", "sha256"])


def update_media_manifest() -> None:
    rows: list[dict[str, object]] = []
    for i in range(1, 7):
        path = UPLOAD / f"Figure_{i:02d}.tiff"
        if path.exists():
            rows.append(
                {
                    "asset_type": "main_figure",
                    "asset": f"Figure_{i:02d}",
                    "upload_file": path.relative_to(OUT).as_posix(),
                    "size_bytes": path.stat().st_size,
                    "sha256": sha256(path),
                }
            )
    for spec in SUPP_FIGS:
        path = UPLOAD / f"{spec.stem}.tiff"
        rows.append(
            {
                "asset_type": "supplementary_figure",
                "asset": spec.stem,
                "upload_file": path.relative_to(OUT).as_posix(),
                "size_bytes": path.stat().st_size,
                "sha256": sha256(path),
            }
        )
    write_csv(QA / "frontiers_upload_figure_manifest_sha256.csv", rows, ["asset_type", "asset", "upload_file", "size_bytes", "sha256"])
    write_csv(QA / "frontiers_upload_media_manifest_sha256.csv", rows, ["asset_type", "asset", "upload_file", "size_bytes", "sha256"])


def make_supp_contact_sheet() -> Path:
    previews = [VECTOR / f"{spec.stem}_preview.png" for spec in SUPP_FIGS]
    thumbs = []
    for idx, path in enumerate(previews, start=1):
        img = Image.open(path).convert("RGB")
        img.thumbnail((760, 560))
        canvas = Image.new("RGB", (820, 640), "white")
        canvas.paste(img, ((820 - img.width) // 2, 28))
        draw = ImageDraw.Draw(canvas)
        draw.text((24, 606), f"Supplementary Figure S{idx}", fill=(48, 60, 72))
        thumbs.append(canvas)
    out = QA / "supplementary_figure_preview_contact_sheet.png"
    sheet = Image.new("RGB", (820, 1280), "white")
    for idx, tile in enumerate(thumbs):
        sheet.paste(tile, (0, idx * 640))
    sheet.save(out, dpi=(150, 150))
    return out


def docx_structural_qa(docx_path: Path) -> dict[str, object]:
    doc = Document(docx_path)
    text = "\n".join(iter_doc_text(doc))
    legend_counts = {
        "Figure 1": len(re.findall(r"(?m)^Figure 1\.", text)),
        "Figure 2": len(re.findall(r"(?m)^Figure 2\.", text)),
        "Figure 3": len(re.findall(r"(?m)^Figure 3\.", text)),
        "Figure 4": len(re.findall(r"(?m)^Figure 4\.", text)),
        "Figure 5": len(re.findall(r"(?m)^Figure 5\.", text)),
        "Figure 6": len(re.findall(r"(?m)^Figure 6\.", text)),
        "Supplementary Figure S1": len(re.findall(r"(?m)^Supplementary Figure S1\.", text)),
        "Supplementary Figure S2": len(re.findall(r"(?m)^Supplementary Figure S2\.", text)),
    }
    zip_names: set[str] = set()
    with zipfile.ZipFile(docx_path) as zf:
        zip_names = set(zf.namelist())
        xml_text = "\n".join(
            zf.read(name).decode("utf-8", errors="ignore")
            for name in zip_names
            if name.startswith("word/") and name.endswith(".xml")
        )
    tracked_change_tags = sum(
        len(re.findall(pattern, xml_text))
        for pattern in (r"<w:ins[\s>]", r"<w:del[\s>]", r"<w:moveFrom[\s>]", r"<w:moveTo[\s>]")
    )
    comments_parts = sorted(name for name in zip_names if name.startswith("word/comments"))
    forbidden_scan = {}
    for term in ["validated endotype", "validated clinical endotype", "clinical prediction model", "causal mechanism"]:
        hits = [line_text.strip() for line_text in text.splitlines() if term.lower() in line_text.lower()]
        forbidden_scan[term] = {
            "count": len(hits),
            "boundary_context_only": all(
                any(marker in hit.lower() for marker in ["excluded", "not ", "from excluded", "do not establish", "rather than"])
                for hit in hits
            ),
            "examples": hits[:3],
        }
    return {
        "docx": rel(docx_path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "legend_counts": legend_counts,
        "old_monocyte_s1_residual": "Supplementary Figure S1. Monocyte pseudobulk context" in text,
        "supplementary_material_pointer_ok": "Supplementary Figures S1-S2 and Supplementary Tables S1 to S22" in text,
        "tracked_change_tags": tracked_change_tags,
        "comments_parts": comments_parts,
        "forbidden_claim_scan": forbidden_scan,
    }


def try_render_docx(docx_path: Path) -> dict[str, object]:
    result: dict[str, object] = {"attempted": False, "success": False, "docx": rel(docx_path)}
    if not DOCX_RENDERER.exists():
        return {**result, "error": "render_docx.py missing"}
    env = os.environ.copy()
    if SOFFICE_APP.exists():
        env["PATH"] = f"{SOFFICE_APP}:{env.get('PATH', '')}"
    out_dir = QA / "manuscript_render_v18"
    if out_dir.exists():
        shutil.rmtree(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    proc = run(
        [
            sys.executable,
            str(DOCX_RENDERER),
            str(docx_path),
            "--output_dir",
            str(out_dir),
            "--emit_pdf",
        ],
        check=False,
        env=env,
    )
    result.update(
        {
            "attempted": True,
            "returncode": proc.returncode,
            "stdout_tail": proc.stdout[-1200:],
            "stderr_tail": proc.stderr[-1200:],
        }
    )
    pdfs = sorted(out_dir.glob("*.pdf"))
    pngs = sorted(out_dir.glob("*.png"))
    result["rendered_png_count"] = len(pngs)
    if proc.returncode == 0 and pdfs and pngs:
        pdf_out = UPLOAD / f"{docx_path.stem}.pdf"
        shutil.copy2(pdfs[0], pdf_out)
        result.update({"success": True, "pdf": rel(pdf_out), "png_dir": rel(out_dir)})
    else:
        result["error"] = "DOCX render did not produce PDF and page PNGs"
    return result


def write_package_manifest() -> None:
    rows: list[dict[str, object]] = []
    for path in sorted(p for p in OUT.rglob("*") if p.is_file()):
        rows.append(
            {
                "relative_path": path.relative_to(OUT).as_posix(),
                "size_bytes": path.stat().st_size,
                "sha256": sha256(path),
            }
        )
    write_csv(QA / "FI_submission_ready_manifest_sha256.csv", rows, ["relative_path", "size_bytes", "sha256"])


def rewrite_readme(docx_render: dict[str, object]) -> None:
    lines = [
        "# FI submission-ready package",
        "",
        f"Generated: {datetime.now().isoformat(timespec='seconds')}",
        "Target journal: Frontiers in Immunology.",
        "",
        "Authoritative main-figure source: user-approved SVG set in the desktop FI figure folder.",
        "This package adds a clean figure-aligned manuscript DOCX and two audit-style supplementary figures. Scientific analyses were not rerun.",
        "",
        "## Upload-ready files",
        "- `01_upload_ready/Frontiers_Immunology_V18_FI_figure_aligned.docx`",
        "- `01_upload_ready/Figure_01.tiff` through `Figure_06.tiff`",
        "- `01_upload_ready/Supplementary_Figure_S1.tiff` and `Supplementary_Figure_S2.tiff`",
    ]
    if docx_render.get("success"):
        lines.append("- `01_upload_ready/Frontiers_Immunology_V18_FI_figure_aligned.pdf`")
    else:
        lines.append("- Manuscript PDF render was not produced; see QA JSON for renderer details.")
    lines.extend(
        [
            "",
            "## Backup and audit",
            "- `02_vector_backup/`: source SVG, outlined SVG, PDF, EPS, PNG and TIFF backups for main and supplementary figures.",
            "- `03_source_data_and_attribution/`: source/audit table copies and attribution for figures.",
            "- `04_QA_and_manifests/`: DOCX structural QA, figure export QA, checksums, contact sheets and zip test.",
            "",
            "## Claim boundary",
            "The manuscript and supplementary figures preserve association-only, artifact-aware interpretation. They do not add clinical prediction, causal mechanism, structural interaction or therapeutic claims.",
        ]
    )
    (OUT / "README.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def zip_submission() -> dict[str, object]:
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(p for p in OUT.rglob("*") if p.is_file()):
            zf.write(path, path.relative_to(ROOT).as_posix())
    proc = run(["unzip", "-t", str(ZIP_PATH)], check=False)
    zip_test = QA / "zip_test.txt"
    zip_test.write_text(proc.stdout + proc.stderr, encoding="utf-8")
    return {
        "zip": rel(ZIP_PATH),
        "returncode": proc.returncode,
        "success": proc.returncode == 0,
        "size_bytes": ZIP_PATH.stat().st_size,
        "sha256": sha256(ZIP_PATH),
        "zip_test": rel(zip_test),
    }


def main() -> None:
    for directory in (UPLOAD, VECTOR, SOURCE, QA):
        directory.mkdir(parents=True, exist_ok=True)
    if not ORIGINAL_DOCX.exists():
        raise FileNotFoundError(ORIGINAL_DOCX)

    docx_edit = revise_docx()
    vector_rows, supp_qa_rows = export_supplementary_figures()
    supp_source_rows = copy_supplementary_source_data()
    update_source_manifest(supp_source_rows)

    write_csv(QA / "supplementary_figure_export_manifest.csv", vector_rows, ["asset", "role", "path", "size_bytes", "sha256"])
    write_csv(QA / "supplementary_figure_QA_summary.csv", supp_qa_rows)
    supp_contact_sheet = make_supp_contact_sheet()

    docx_qa = docx_structural_qa(REVISED_DOCX)
    render = try_render_docx(REVISED_DOCX)
    update_media_manifest()
    rewrite_readme(render)

    final_report_path = QA / "final_QA_report.json"
    existing_final = {}
    if final_report_path.exists():
        try:
            existing_final = json.loads(final_report_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            existing_final = {"previous_final_QA_report_parse_error": True}

    report = {
        **existing_final,
        "updated_at": datetime.now().isoformat(timespec="seconds"),
        "target_journal": "Frontiers in Immunology",
        "frontiers_requirements_checked": {
            "main_figures": "TIF/TIFF or JPEG; RGB; 300 dpi at final size",
            "supplementary_images": "CDX/EPS/JPEG/PDF/PNG/TIF/TIFF; 300 dpi; RGB",
            "source_urls": [
                "https://www.frontiersin.org/journals/immunology/for-authors/author-guidelines",
                "https://www.frontiersin.org/journals/immunology/for-authors/submission-checklist",
            ],
        },
        "docx_edit": docx_edit,
        "docx_structural_QA": docx_qa,
        "docx_render_v18": render,
        "supplementary_figures": supp_qa_rows,
        "supplementary_contact_sheet": rel(supp_contact_sheet),
        "external_vector_assets_used": False,
        "frontiers_upload_checklist_update": {
            "six_main_tiff_files_present": all((UPLOAD / f"Figure_{i:02d}.tiff").exists() for i in range(1, 7)),
            "two_supplementary_tiff_files_present": all((UPLOAD / f"{spec.stem}.tiff").exists() for spec in SUPP_FIGS),
            "supplementary_tiffs_rgb": all(row["tiff_mode"] == "RGB" for row in supp_qa_rows),
            "supplementary_tiffs_at_least_300_dpi": all(row["frontiers_min_dpi_pass"] for row in supp_qa_rows),
            "supplementary_svgs_have_no_embedded_rasters": all(row["embedded_raster_count"] == 0 for row in supp_qa_rows),
            "supplementary_outlined_svgs_have_no_live_text": all(row["outlined_text_to_path_pass"] for row in supp_qa_rows),
            "docx_no_old_monocyte_s1_legend": not docx_qa["old_monocyte_s1_residual"],
        },
    }
    final_report_path.write_text(json.dumps(report, indent=2), encoding="utf-8")
    write_csv(QA / "manuscript_text_QA_summary.csv", [
        {
            "check": "Figure 1-6 legends each exactly one",
            "pass": all(docx_qa["legend_counts"][f"Figure {i}"] == 1 for i in range(1, 7)),
            "details": json.dumps(docx_qa["legend_counts"], ensure_ascii=False),
        },
        {
            "check": "Supplementary Figure S1-S2 legends each exactly one",
            "pass": docx_qa["legend_counts"]["Supplementary Figure S1"] == 1 and docx_qa["legend_counts"]["Supplementary Figure S2"] == 1,
            "details": json.dumps(docx_qa["legend_counts"], ensure_ascii=False),
        },
        {
            "check": "Old monocyte pseudobulk S1 legend removed",
            "pass": not docx_qa["old_monocyte_s1_residual"],
            "details": "",
        },
        {
            "check": "No tracked changes or comments",
            "pass": docx_qa["tracked_change_tags"] == 0 and not docx_qa["comments_parts"],
            "details": json.dumps({"tracked_change_tags": docx_qa["tracked_change_tags"], "comments_parts": docx_qa["comments_parts"]}, ensure_ascii=False),
        },
    ])

    write_package_manifest()
    zip_result = zip_submission()
    if not zip_result["success"]:
        raise RuntimeError("Final zip failed unzip -t")

    print(json.dumps(
        {
            "revised_docx": str(REVISED_DOCX),
            "output_dir": str(OUT),
            "zip": str(ZIP_PATH),
            "supplementary_figures": [row["supplementary_figure"] for row in supp_qa_rows],
            "docx_render_success": render.get("success", False),
            "zip_test_success": zip_result["success"],
            "supplementary_contact_sheet": str(supp_contact_sheet),
        },
        indent=2,
    ))


if __name__ == "__main__":
    main()
