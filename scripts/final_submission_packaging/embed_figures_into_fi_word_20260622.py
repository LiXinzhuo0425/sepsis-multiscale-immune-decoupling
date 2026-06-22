#!/usr/bin/env python3
"""Embed finalized FI figures into the two final Word manuscripts.

Positions:
- Main Figures 1-6 are inserted immediately after their corresponding figure
  legend paragraphs in both Word variants.
- Supplementary Figures S1-S2 are inserted immediately after their corresponding
  supplementary figure legend paragraphs in the with-supplementary variant.

The source manuscripts and separate upload-ready TIFF files are not modified.
"""

from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "06_FI_SUBMISSION_READY_20260620"
UPLOAD = OUT / "01_upload_ready"
VECTOR = OUT / "02_vector_backup"
QA = OUT / "04_QA_and_manifests"
ZIP_PATH = ROOT / "FI_submission_ready_20260620.zip"

WITH_SUPP_SOURCE = ROOT / "00_FINAL_WORD/Frontiers_Immunology_V19_FI_SUBMISSION_WITH_SUPPLEMENTARY.docx"
MAIN_ONLY_SOURCE = ROOT / "00_FINAL_WORD/Frontiers_Immunology_V19_FI_SUBMISSION_MAIN_TEXT_ONLY.docx"
WITH_SUPP_OUT = ROOT / "00_FINAL_WORD/Frontiers_Immunology_V20_FI_SUBMISSION_WITH_FIGURES_AND_SUPPLEMENTARY.docx"
MAIN_ONLY_OUT = ROOT / "00_FINAL_WORD/Frontiers_Immunology_V20_FI_SUBMISSION_MAIN_TEXT_WITH_FIGURES.docx"

DOCX_RENDERER = Path(
    "<CODEX_HOME>/plugins/cache/openai-primary-runtime/documents/26.619.11828/skills/documents/render_docx.py"
)
SOFFICE_APP = Path("/Applications/LibreOffice.app/Contents/MacOS")


MAIN_FIGURES = [
    ("Figure 1.", VECTOR / "Figure_01_600dpi.png", 6.25),
    ("Figure 2.", VECTOR / "Figure_02_600dpi.png", 6.25),
    ("Figure 3.", VECTOR / "Figure_03_600dpi.png", 6.25),
    ("Figure 4.", VECTOR / "Figure_04_600dpi.png", 6.25),
    ("Figure 5.", VECTOR / "Figure_05_600dpi.png", 5.65),
    ("Figure 6.", VECTOR / "Figure_06_600dpi.png", 6.25),
]
SUPP_FIGURES = [
    ("Supplementary Figure S1.", VECTOR / "Supplementary_Figure_S1_600dpi.png", 6.25),
    ("Supplementary Figure S2.", VECTOR / "Supplementary_Figure_S2_600dpi.png", 6.25),
]


def run(cmd: list[str], *, check: bool = True, env: dict[str, str] | None = None) -> subprocess.CompletedProcess[str]:
    proc = subprocess.run(
        cmd,
        cwd=str(ROOT),
        env=env,
        check=False,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    if check and proc.returncode != 0:
        raise RuntimeError(
            "Command failed:\n"
            + " ".join(cmd)
            + f"\nexit={proc.returncode}\nSTDOUT:\n{proc.stdout}\nSTDERR:\n{proc.stderr}"
        )
    return proc


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def rel(path: Path) -> str:
    try:
        return path.relative_to(ROOT).as_posix()
    except ValueError:
        return str(path)


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_caption_paragraph(doc: Document, prefix: str) -> Paragraph:
    matches = [p for p in doc.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected exactly one caption for {prefix}; found {len(matches)}")
    return matches[0]


def format_caption_for_image(caption: Paragraph, page_break_before: bool) -> None:
    pf = caption.paragraph_format
    pf.keep_with_next = True
    pf.keep_together = False
    pf.page_break_before = page_break_before
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.0


def insert_image_after_caption(caption: Paragraph, image_path: Path, width_in: float) -> None:
    image_para = insert_paragraph_after(caption)
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = image_para.paragraph_format
    pf.keep_together = True
    pf.keep_with_next = False
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    run_obj = image_para.add_run()
    run_obj.add_picture(str(image_path), width=Inches(width_in))


def paragraph_text(element) -> str:
    return "".join(t.text or "" for t in element.iter(qn("w:t")))


def child_tag(element) -> str:
    return element.tag.rsplit("}", 1)[-1]


def set_page_break_before_heading(doc: Document, heading_text: str) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == heading_text:
            paragraph.paragraph_format.page_break_before = True
            return


def embed_for_doc(docx_in: Path, docx_out: Path, include_supp: bool) -> dict[str, object]:
    doc = Document(docx_in)
    inserted: list[dict[str, object]] = []

    for index, (prefix, image_path, width_in) in enumerate(MAIN_FIGURES):
        if not image_path.exists():
            raise FileNotFoundError(image_path)
        caption = find_caption_paragraph(doc, prefix)
        format_caption_for_image(caption, page_break_before=index > 0)
        insert_image_after_caption(caption, image_path, width_in)
        inserted.append({"caption": prefix, "image": rel(image_path), "width_in": width_in})

    if include_supp:
        for index, (prefix, image_path, width_in) in enumerate(SUPP_FIGURES):
            if not image_path.exists():
                raise FileNotFoundError(image_path)
            caption = find_caption_paragraph(doc, prefix)
            format_caption_for_image(caption, page_break_before=index > 0)
            insert_image_after_caption(caption, image_path, width_in)
            inserted.append({"caption": prefix, "image": rel(image_path), "width_in": width_in})

    set_page_break_before_heading(doc, "Tables")
    if include_supp:
        set_page_break_before_heading(doc, "Supplementary Materials")

    doc.save(docx_out)
    shutil.copy2(docx_out, UPLOAD / docx_out.name)
    return {"source": rel(docx_in), "output": rel(docx_out), "inserted": inserted}


def iter_text(doc: Document) -> Iterable[str]:
    for p in doc.paragraphs:
        yield p.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p.text


def count_embedded_drawings(docx_path: Path) -> dict[str, object]:
    with zipfile.ZipFile(docx_path) as zf:
        names = zf.namelist()
        doc_xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        media = sorted(name for name in names if name.startswith("word/media/"))
        xml_all = "\n".join(
            zf.read(name).decode("utf-8", errors="ignore")
            for name in names
            if name.startswith("word/") and name.endswith(".xml")
        )
    document_root = ET.fromstring(doc_xml.encode("utf-8"))
    local_names = [element.tag.rsplit("}", 1)[-1] for element in document_root.iter()]
    drawing_elements = sum(1 for name in local_names if name == "drawing")
    inline_or_anchor_elements = sum(1 for name in local_names if name in {"inline", "anchor"})
    tracked_change_tags = sum(
        len(re.findall(pattern, xml_all))
        for pattern in (r"<w:ins[\s>]", r"<w:del[\s>]", r"<w:moveFrom[\s>]", r"<w:moveTo[\s>]")
    )
    comments_parts = sorted(name for name in names if name.startswith("word/comments"))
    return {
        "drawing_elements": drawing_elements,
        "inline_or_anchor_elements": inline_or_anchor_elements,
        "media_files": media,
        "media_count": len(media),
        "tracked_change_tags": tracked_change_tags,
        "comments_parts": comments_parts,
    }


def structural_qa(docx_path: Path, expected_main: int, expected_supp: int) -> dict[str, object]:
    doc = Document(docx_path)
    text = "\n".join(iter_text(doc))
    counts = count_embedded_drawings(docx_path)
    legend_counts = {f"Figure {i}": len(re.findall(rf"(?m)^Figure {i}\.", text)) for i in range(1, 7)}
    legend_counts.update(
        {
            "Supplementary Figure S1": len(re.findall(r"(?m)^Supplementary Figure S1\.", text)),
            "Supplementary Figure S2": len(re.findall(r"(?m)^Supplementary Figure S2\.", text)),
        }
    )
    return {
        "docx": rel(docx_path),
        "sha256": sha256(docx_path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "legend_counts": legend_counts,
        "expected_embedded_figures": expected_main + expected_supp,
        **counts,
        "embedded_figure_count_pass": (
            counts["media_count"] == expected_main + expected_supp
            and counts["drawing_elements"] >= expected_main + expected_supp
            and counts["inline_or_anchor_elements"] >= expected_main + expected_supp
        ),
        "comments_pass": not counts["comments_parts"],
        "tracked_changes_pass": counts["tracked_change_tags"] == 0,
    }


def render_docx(docx_path: Path, variant: str) -> dict[str, object]:
    out_dir = QA / f"render_{variant}"
    if out_dir.exists():
        shutil.rmtree(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    env = os.environ.copy()
    if SOFFICE_APP.exists():
        env["PATH"] = f"{SOFFICE_APP}:{env.get('PATH', '')}"
    proc = run(
        [sys.executable, str(DOCX_RENDERER), str(docx_path), "--output_dir", str(out_dir), "--emit_pdf"],
        check=False,
        env=env,
    )
    pdfs = sorted(out_dir.glob("*.pdf"))
    pngs = sorted(out_dir.glob("page-*.png"))
    result = {
        "variant": variant,
        "attempted": True,
        "success": proc.returncode == 0 and bool(pdfs) and bool(pngs),
        "returncode": proc.returncode,
        "rendered_png_count": len(pngs),
        "stdout_tail": proc.stdout[-1200:],
        "stderr_tail": proc.stderr[-1200:],
        "png_dir": rel(out_dir),
    }
    if result["success"]:
        pdf_out = UPLOAD / f"{docx_path.stem}.pdf"
        shutil.copy2(pdfs[0], pdf_out)
        result["pdf"] = rel(pdf_out)
    return result


def make_contact_sheet(render_dir: Path, out_path: Path) -> None:
    pages = sorted(render_dir.glob("page-*.png"), key=lambda p: int(p.stem.split("-")[1]))
    thumbs = []
    for page in pages:
        img = Image.open(page).convert("RGB")
        img.thumbnail((250, 330))
        tile = Image.new("RGB", (290, 370), "white")
        tile.paste(img, ((290 - img.width) // 2, 8))
        thumbs.append((tile, page.stem))
    cols = 4
    rows = (len(thumbs) + cols - 1) // cols
    sheet = Image.new("RGB", (cols * 290, rows * 370), "white")
    from PIL import ImageDraw

    draw = ImageDraw.Draw(sheet)
    for idx, (tile, label) in enumerate(thumbs):
        x = (idx % cols) * 290
        y = (idx // cols) * 370
        sheet.paste(tile, (x, y))
        draw.text((x + 8, y + 346), label, fill=(0, 0, 0))
    sheet.save(out_path, dpi=(150, 150))


def write_csv(path: Path, rows: list[dict[str, object]]) -> None:
    keys: list[str] = []
    for row in rows:
        for key in row:
            if key not in keys:
                keys.append(key)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=keys)
        writer.writeheader()
        writer.writerows(rows)


def write_package_manifest() -> None:
    rows = []
    for path in sorted(p for p in OUT.rglob("*") if p.is_file()):
        rows.append({"relative_path": path.relative_to(OUT).as_posix(), "size_bytes": path.stat().st_size, "sha256": sha256(path)})
    write_csv(QA / "FI_submission_ready_manifest_sha256.csv", rows)


def update_readme() -> None:
    readme = OUT / "README.md"
    current = readme.read_text(encoding="utf-8") if readme.exists() else "# FI submission-ready package\n"
    marker = "\n## Embedded-figure Word variants\n"
    current = current.split(marker, 1)[0].rstrip()
    addition = f"""

## Embedded-figure Word variants
- `01_upload_ready/{WITH_SUPP_OUT.name}`: manuscript with Figure 1-6 plus Supplementary Figure S1-S2 embedded after their legends.
- `01_upload_ready/{MAIN_ONLY_OUT.name}`: main manuscript with Figure 1-6 embedded after their legends, without the appended supplementary-material appendix.
- Separate TIFF figure files remain in `01_upload_ready/` for journal upload.
"""
    readme.write_text(current + addition + "\n", encoding="utf-8")


def zip_submission() -> dict[str, object]:
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    zip_test_path = QA / "zip_test.txt"
    if zip_test_path.exists():
        zip_test_path.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(p for p in OUT.rglob("*") if p.is_file()):
            zf.write(path, path.relative_to(ROOT).as_posix())
    proc = run(["unzip", "-t", str(ZIP_PATH)], check=False)
    zip_test_path.write_text(proc.stdout + proc.stderr, encoding="utf-8")
    return {"zip": rel(ZIP_PATH), "success": proc.returncode == 0, "returncode": proc.returncode, "sha256": sha256(ZIP_PATH), "size_bytes": ZIP_PATH.stat().st_size}


def clean_upload_intermediate_word_versions() -> None:
    for stale in UPLOAD.glob("Frontiers_Immunology_V19_FI_SUBMISSION_*.docx"):
        stale.unlink()
    for stale in UPLOAD.glob("Frontiers_Immunology_V19_FI_SUBMISSION_*.pdf"):
        stale.unlink()


def main() -> None:
    for path in [WITH_SUPP_SOURCE, MAIN_ONLY_SOURCE, DOCX_RENDERER]:
        if not path.exists():
            raise FileNotFoundError(path)
    for _, image_path, _ in MAIN_FIGURES + SUPP_FIGURES:
        if not image_path.exists():
            raise FileNotFoundError(image_path)
    QA.mkdir(parents=True, exist_ok=True)
    UPLOAD.mkdir(parents=True, exist_ok=True)

    clean_upload_intermediate_word_versions()
    created = [
        embed_for_doc(WITH_SUPP_SOURCE, WITH_SUPP_OUT, include_supp=True),
        embed_for_doc(MAIN_ONLY_SOURCE, MAIN_ONLY_OUT, include_supp=False),
    ]
    qa = [
        structural_qa(WITH_SUPP_OUT, expected_main=6, expected_supp=2),
        structural_qa(MAIN_ONLY_OUT, expected_main=6, expected_supp=0),
    ]
    render = [
        render_docx(WITH_SUPP_OUT, "v20_with_figures_and_supplementary"),
        render_docx(MAIN_ONLY_OUT, "v20_main_text_with_figures"),
    ]
    for item in render:
        if item["success"]:
            make_contact_sheet(
                ROOT / item["png_dir"],
                QA / f"{item['variant']}_contact_sheet.png",
            )

    update_readme()
    report = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "created": created,
        "structural_QA": qa,
        "render_QA": render,
        "zip_result_note": "Final zip test is written after packaging to 04_QA_and_manifests/zip_test.txt and reported in the console output.",
    }
    (QA / "embedded_figure_word_QA_report.json").write_text(json.dumps(report, indent=2), encoding="utf-8")
    write_package_manifest()
    zip_result = zip_submission()
    if not zip_result["success"]:
        raise RuntimeError("Final zip failed unzip -t")
    report["zip_result"] = zip_result
    (QA / "embedded_figure_word_QA_report.json").write_text(json.dumps(report, indent=2), encoding="utf-8")
    print(json.dumps(
        {
            "with_supplementary_embedded": str(WITH_SUPP_OUT),
            "main_only_embedded": str(MAIN_ONLY_OUT),
            "render_success": [item["success"] for item in render],
            "zip_success": zip_result["success"],
            "zip": str(ZIP_PATH),
        },
        indent=2,
    ))


if __name__ == "__main__":
    main()
