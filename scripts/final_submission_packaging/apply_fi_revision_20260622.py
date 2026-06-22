#!/usr/bin/env python3
"""Apply the final FI revision requested on 2026-06-22.

This script is a packaging and manuscript-editing pass only. It does not rerun
scientific analyses. It creates a clean V21 submission package
with:
- a revised main manuscript containing text and figure legends only;
- a separate narrow supplementary-materials index document;
- upload-ready TIFF figures copied from the approved package, with explicit
  overrides for figure text that was revised after the approved package;
- editable CSV supplementary tables S1-S22 in a zip archive;
- render, structural, and packaging QA.
"""

from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw


ROOT = Path(__file__).resolve().parents[1]
OLD_OUT = ROOT / "06_FI_SUBMISSION_READY_20260620"
OLD_UPLOAD = OLD_OUT / "01_upload_ready"
OUT = ROOT / "07_FI_SUBMISSION_REVISED_20260622"
UPLOAD = OUT / "01_upload_ready"
QA = OUT / "04_QA_and_manifests"
TABLE_DIR = UPLOAD / "Supplementary_Tables_S1_S22_CSV"
ZIP_PATH = ROOT / "FI_submission_revised_20260622.zip"

SOURCE_MAIN = ROOT / "00_FINAL_WORD/Frontiers_Immunology_V19_FI_SUBMISSION_MAIN_TEXT_ONLY.docx"
MAIN_DOCX = ROOT / "00_FINAL_WORD/Frontiers_Immunology_V21_FI_SUBMISSION_MAIN_TEXT_ONLY_REVISED.docx"
SUPP_DOCX = ROOT / "00_FINAL_WORD/Frontiers_Immunology_V21_FI_SUPPLEMENTARY_MATERIALS_INDEX_ONLY.docx"

DOCX_RENDERER = Path(
    "<CODEX_DOCUMENTS_RENDER_DOCX>"
)
SOFFICE_APP = Path("/Applications/LibreOffice.app/Contents/MacOS")

SUPP_INDEX = (
    ROOT
    / "03_TRACEABLE_RESULTS/stage_outputs_tables/03_REPRODUCIBLE_RESULTS/"
    "stage7_review_hardening/stage7_supplementary_table_index.csv"
)

FIGURE_FILES = [f"Figure_{i:02d}.tiff" for i in range(1, 7)] + [
    "Supplementary_Figure_S1.tiff",
    "Supplementary_Figure_S2.tiff",
]
PREVIEW_FILES = [f"Figure_{i:02d}_preview.png" for i in range(1, 7)] + [
    "Supplementary_Figure_S1_preview.png",
    "Supplementary_Figure_S2_preview.png",
]

FIGURE_FILE_OVERRIDES = {
    "Figure_01.tiff": ROOT / "figures/Figure1_study_architecture_claim_boundary.live-text_600dpi.tiff",
    "Figure_01_preview.png": ROOT / "figures/Figure1_study_architecture_claim_boundary.live-text_preview.png",
}


TITLE = (
    "Compartment-aware reconstruction of inflammatory/MHC-II decoupling in sepsis "
    "using public RNA sequencing and cytometry data"
)

CONTRIBUTION = (
    "Sepsis often combines inflammatory activation with impaired antigen presentation, but RNA profiling "
    "and cytometry measure different immune layers. This study asks how those layers relate using public "
    "datasets rather than new samples. Across seven public sepsis RNA cohorts, inflammatory programs "
    "increased while antigen-presentation programs involving MHC-II, CD74 and HLA-DR decreased or "
    "inversely coupled with inflammation. In COMBAT, paired RNA and mass cytometry measurements linked "
    "this RNA pattern to a representative residual co-event signal within a broader marker-pair co-event "
    "background. Raw-FCS analyses showed that the prespecified CD3/CD14 metric was not uniquely specific, "
    "so the cytometry result is interpreted as an artifact-aware event-derived summary rather than a "
    "structural interaction or new cell subset. Case-only analyses supported a continuous immune-decoupling "
    "gradient, and donor-level single-cell and pathway analyses supplied context. The contribution is a "
    "reproducible, claim-bounded framework for integrating public RNA, cytometry and single-cell layers "
    "in sepsis. By separating recurrence, artifact controls and claim boundaries, it clarifies what public "
    "multi-layer data can and cannot support. It is association-only and requires prospective paired RNA, "
    "cytometry and protein-level validation before clinical or mechanistic use."
)

ABSTRACT = (
    "Background: Whole-blood RNA sequencing and cytometry assays capture different immune analytical "
    "layers, which can obscure how inflammation and immunoparalysis coexist in sepsis. Methods: We "
    "integrated seven public bulk sepsis transcriptomic cohorts, derived cytometry summary tables, COMBAT "
    "paired RNA sequencing and mass cytometry participant-timepoint data, matched raw-FCS sensitivity "
    "analyses, and donor-level single-cell contextual analysis. Results: A total of 660 bulk RNA samples "
    "(499 sepsis cases) were included. The prespecified six-gene host-response panel showed inverse "
    "coupling with MHC-II/CD74 across all seven cohorts (median Spearman rho approximately -0.60), and "
    "composition sensitivity attenuated but did not reverse the primary direction. Case-only analysis "
    "supported a continuous immune-decoupling gradient rather than a validated clinical endotype. In "
    "COMBAT sepsis-only paired samples (40 rows from 34 participants), a representative residual co-event "
    "signal within a broader marker-pair co-event background, operationalized by the prespecified CD3/CD14 "
    "metric, was negatively associated with HLA-DR core RNA score (rho = -0.559, permutation p = 4.00 x "
    "10^-4) and MHC-II/CD74 RNA score (rho = -0.517, permutation p = 0.0011), and positively associated "
    "with RNA decoupling index (rho = 0.462, permutation p = 0.0035). Row-level, participant-level, "
    "participant-cluster, and composition-adjusted sensitivities preserved association directions with "
    "expected attenuation. Raw-FCS all-pair analysis did not resolve marker-pair specificity; therefore "
    "the cytometry finding is interpreted as representative and artifact-aware rather than CD3/CD14-specific. "
    "Conclusion: These findings support a bounded systems immunology reconstruction of sepsis-associated "
    "immunoparalysis that requires prospective paired biological validation."
)


PARAGRAPH_REPLACEMENTS = {
    "Study design and target format": "Study design and scope",
    (
        "This was a public-data systems immunology study designed for submission to Frontiers in Immunology, "
        "Systems Immunology section, as an Original Research article. No MIMIC-IV, eICU, protected health "
        "information, new human samples, wet-laboratory experiments, MR/SMR causal analysis, clinical "
        "prediction modeling, drug repositioning, or treatment-response modeling was performed."
    ): (
        "This was a public-data systems immunology secondary computational analysis. No MIMIC-IV, eICU, "
        "protected health information, new human samples, wet-laboratory experiments, MR/SMR causal analysis, "
        "clinical prediction modeling, drug repositioning, or treatment-response modeling was performed."
    ),
    (
        "Here we investigated whether public RNA and cytometry data could support a bounded, compartment-aware "
        "model of sepsis-associated immunoparalysis. The study architecture and claim boundaries are summarized "
        "in Figure 1. Our evidence hierarchy proceeds from cross-cohort bulk RNA recurrence to a continuous "
        "case-only immune-decoupling spectrum, raw-FCS pair-specificity hardening of the residual CD3/CD14 "
        "co-event signal, a paired COMBAT RNA-CyTOF bridge, donor-level single-cell contextual support, pathway "
        "context and an exploratory 28-day mortality clinical anchor. The central conclusion is constrained to "
        "the following: public RNA and cytometry data support a compartment-aware model of sepsis-associated "
        "immunoparalysis, in which systemic inflammatory/MHC-II decoupling is linked to a prespecified residual "
        "CD3/CD14 co-event signal in paired sepsis samples, while the event-derived signal itself remains an "
        "artifact-aware summary rather than a specific structural interaction or validated cell subset."
    ): (
        "Here we investigated whether public RNA and cytometry data could support a bounded, compartment-aware "
        "model of sepsis-associated immunoparalysis. The study architecture and claim boundaries are summarized "
        "in Figure 1. Our evidence hierarchy proceeds from cross-cohort bulk RNA recurrence to a continuous "
        "case-only immune-decoupling spectrum, raw-FCS marker-pair sensitivity around a representative residual "
        "co-event signal, a paired COMBAT RNA-CyTOF bridge, donor-level single-cell contextual support, pathway "
        "context and an exploratory 28-day mortality clinical anchor. The central conclusion is constrained to "
        "the following: public RNA and cytometry data support a compartment-aware model of sepsis-associated "
        "immunoparalysis, in which systemic inflammatory/MHC-II decoupling is associated with a representative "
        "residual co-event signal within a broader marker-pair co-event background, while the event-derived "
        "signal itself remains an artifact-aware summary rather than a specific structural interaction or "
        "validated cell subset."
    ),
    "Residual CD3/CD14 co-event signals showed state remodeling rather than uniform HLA-DR loss": (
        "Residual co-event summaries showed marker-state heterogeneity rather than uniform HLA-DR loss"
    ),
    "COMBAT paired RNA-CyTOF linked CD3/CD14 co-event signal to RNA immunoparalysis in Sepsis-only samples": (
        "COMBAT paired RNA-CyTOF linked a representative residual co-event signal to RNA antigen-presentation context in sepsis-only samples"
    ),
    (
        "All analyses used public de-identified resources or local pre-analyzed derived summaries; no new ethics "
        "approval was required for the present secondary computational analysis."
    ): (
        "All analyses used public de-identified resources or derived aggregate summaries generated from public "
        "datasets and deposited with the analysis outputs; no new ethics approval was required for the present "
        "secondary computational analysis."
    ),
    (
        "Cytometry summaries were interpreted with explicit artifact boundaries. Frequency and summary-level "
        "marker-state context were treated as distinct readouts. Matching public raw whole-blood FCS files were "
        "available for all 40 COMBAT sepsis-only paired rows and enabled event-level negative-pair and "
        "pair-specificity hardening. In the required marker-pair comparison, CD3/CD14 showed the expected "
        "association direction but was not uniquely strongest: for the RNA decoupling index, CD14/CD16 and "
        "CD4/CD14 showed larger rho values than CD3/CD14, whereas CD19/CD14, CD8/CD14 and CD56/CD14 were weaker "
        "or opposite. In the all-pair raw-FCS marker-null analysis of 378 available marker pairs, CD3/CD14 "
        "ranked in the 61.6% to 65.1% percentile range by absolute rho across the three RNA metrics, and 4 to 5 "
        "of 12 abundance-matched control pairs showed absolute correlations at least as large as CD3/CD14. "
        "Threshold perturbation and technical-covariate residualization preserved the expected direction for "
        "CD3/CD14, whereas event-count downsampling weakened this pattern. These results reduce the earlier "
        "feasibility gap but leave CD3/CD14 pair-level specificity unresolved; CD3/CD14 is therefore retained "
        "as a prespecified representative residual co-event signal, not as evidence for a specific structural "
        "interaction or new cell subset (Figure 4; Supplementary Tables S8-S10 and S19)."
    ): (
        "Cytometry summaries were interpreted with explicit artifact boundaries. Frequency and summary-level "
        "marker-state context were treated as distinct readouts. Matching public raw whole-blood FCS files were "
        "available for all 40 COMBAT sepsis-only paired rows and enabled event-level negative-pair and "
        "pair-specificity sensitivity analyses. In the required marker-pair comparison, the prespecified "
        "CD3/CD14 metric showed the expected association direction but was not uniquely strongest: for the RNA "
        "decoupling index, CD14/CD16 and CD4/CD14 showed larger rho values than CD3/CD14, whereas CD19/CD14, "
        "CD8/CD14 and CD56/CD14 were weaker or opposite. In the all-pair raw-FCS marker-null analysis of 378 "
        "available marker pairs, CD3/CD14 ranked in the 61.6% to 65.1% percentile range by absolute rho across "
        "the three RNA metrics, and 4 to 5 of 12 abundance-matched control pairs showed absolute correlations "
        "at least as large as CD3/CD14. Threshold perturbation and technical-covariate residualization preserved "
        "the expected direction for CD3/CD14, whereas event-count downsampling weakened this pattern. These "
        "results reduce the earlier feasibility gap but leave pair-level specificity unresolved; the cytometry "
        "readout is therefore retained as a representative residual co-event signal within a broader marker-pair "
        "co-event background, not as evidence for a specific structural interaction or new cell subset (Figure 4; "
        "Supplementary Tables S8-S10 and S19)."
    ),
    (
        "Within sepsis residual CD3/CD14 events, summary-level marker-state context included HLA-DR positivity "
        "(median 0.100, IQR 0.029-0.223), CD33 positivity (median 0.073, IQR 0.014-0.142), CD11c positivity "
        "(median 0.067, IQR 0.009-0.202) and CD38 positivity (median 0.317, IQR 0.186-0.494). These summaries "
        "make the activation-state remodeling more concrete, but they do not define a new cell subset or prove "
        "structural contact."
    ): (
        "Within sepsis residual CD3/CD14 events, summary-level marker-state context included HLA-DR positivity "
        "(median 0.100, IQR 0.029-0.223), CD33 positivity (median 0.073, IQR 0.014-0.142), CD11c positivity "
        "(median 0.067, IQR 0.009-0.202) and CD38 positivity (median 0.317, IQR 0.186-0.494). These summaries "
        "make the activation-state context more concrete, but they do not define a new cell subset or prove "
        "structural contact."
    ),
    (
        "COMBAT paired analysis provided the most direct paired public-data bridge in this study. The paired "
        "dataset contained 129 matched biological RNA-CyTOF participant-timepoint rows, including 40 sepsis-only "
        "rows from 34 participants. Six participants contributed repeated sepsis timepoints, with a maximum of "
        "two rows per participant. In sepsis-only samples, abundance-normalized residual CD3/CD14 co-event signal "
        "was inversely associated with HLA-DR core RNA score and MHC-II/CD74 RNA score, and positively associated "
        "with the RNA decoupling index. Row-level correlations, participant-level averaged correlations, "
        "participant-cluster bootstrap, and participant-cluster rank-regression sensitivity preserved the same "
        "association directions, although uncertainty increased after collapsing or resampling repeated "
        "observations (Figure 5; Table 5; Supplementary Table S16)."
    ): (
        "COMBAT paired analysis provided the most direct paired public-data bridge in this study. The paired "
        "dataset contained 129 matched biological RNA-CyTOF participant-timepoint rows, including 40 sepsis-only "
        "rows from 34 participants. Six participants contributed repeated sepsis timepoints, with a maximum of "
        "two rows per participant. In sepsis-only samples, a representative residual co-event signal within a "
        "broader marker-pair co-event background, operationalized by the prespecified CD3/CD14 metric, was "
        "inversely associated with HLA-DR core RNA score and MHC-II/CD74 RNA score, and positively associated "
        "with the RNA decoupling index. Row-level correlations, participant-level averaged correlations, "
        "participant-cluster bootstrap, and participant-cluster rank-regression sensitivity preserved the same "
        "association directions, although uncertainty increased after collapsing or resampling repeated "
        "observations (Figure 5; Table 5; Supplementary Table S16)."
    ),
    (
        "This study supports a bounded systems immunology reconstruction of sepsis-associated immunoparalysis. "
        "The fundamental contribution is not the demonstration that MHC-II or HLA-DR suppression exists in "
        "sepsis-this is already well established-but rather that systemic inflammatory/MHC-II decoupling can be "
        "linked to a cytometry-defined residual CD3/CD14 co-event signal in matched sepsis samples, while the "
        "residual event-derived summary itself exhibits activation-state remodeling."
    ): (
        "This study supports a bounded systems immunology reconstruction of sepsis-associated immunoparalysis. "
        "The fundamental contribution is not the demonstration that MHC-II or HLA-DR suppression exists in "
        "sepsis, as this is already well established, but rather that systemic inflammatory/MHC-II decoupling can "
        "be associated with a representative residual co-event signal within a broader marker-pair co-event "
        "background in matched sepsis samples, while the residual event-derived summary itself exhibits "
        "activation-state context."
    ),
    (
        "The cytometry evidence also necessitates conservative reformulation. Raw-FCS pair-specificity hardening "
        "did not establish CD3/CD14 as a uniquely specific pair; rather, CD3/CD14 remains a prespecified "
        "representative residual co-event signal within a broader co-event background. The safer interpretation "
        "is one of analytical-layer divergence: whole-blood RNA can capture population-level antigen-presentation "
        "suppression while residual event-derived summaries show summary-level marker-state context and co-event "
        "frequency behavior that are not reducible to a single structural interaction claim."
    ): (
        "The cytometry evidence also necessitates conservative reformulation. Raw-FCS pair-specificity hardening "
        "did not establish CD3/CD14 as a uniquely specific pair; rather, CD3/CD14 remains a representative "
        "residual co-event signal within a broader marker-pair co-event background. The safer interpretation is "
        "one of analytical-layer divergence: whole-blood RNA can capture population-level antigen-presentation "
        "suppression while residual event-derived summaries show summary-level marker-state context and co-event "
        "frequency behavior that are not reducible to a single structural interaction claim."
    ),
    (
        "The stronger CD14/CD16 and CD4/CD14 associations provide a biologically plausible caution against "
        "over-centering a single pair: CD14/CD16 may capture broader myeloid and monocyte-lineage co-event "
        "background, whereas CD4/CD14 may be closer to helper T-cell/monocyte antigen-presentation context. "
        "Accordingly, the study treats CD3/CD14 as a prespecified representative co-event signal rather than as "
        "a uniquely specific pair."
    ): (
        "The stronger CD14/CD16 and CD4/CD14 associations provide a biologically plausible caution against "
        "over-centering a single pair: CD14/CD16 may capture broader myeloid and monocyte-lineage co-event "
        "background, whereas CD4/CD14 may be closer to helper T-cell/monocyte antigen-presentation context. "
        "Accordingly, the study treats CD3/CD14 as a representative residual co-event signal within a broader "
        "marker-pair co-event background rather than as a uniquely specific pair."
    ),
    (
        "The validation strategy reflects the inherent constraints of public-data computational research. "
        "Frontiers-relevant validation is addressed through cross-cohort recurrence, continuous-spectrum analysis, "
        "broad-cell composition sensitivity, orthogonal paired RNA-CyTOF bridging, participant-level robustness, "
        "raw-FCS artifact-control hardening, single-cell donor-level contextual analysis, pathway enrichment "
        "context and explicit claim-boundary auditing. These analyses strengthen the evidence chain but do not "
        "replace the requirement for prospective biological validation. The resulting bounded working "
        "interpretation is summarized in Figure 6, with the evidence hierarchy, claim-boundary matrix and "
        "reproducibility traceability summarized in Supplementary Figures S1 and S2."
    ): (
        "The validation strategy reflects the inherent constraints of public-data computational research. "
        "Validation was addressed through cross-cohort recurrence, continuous-spectrum analysis, composition "
        "sensitivity, paired RNA-CyTOF bridging, participant-level robustness, raw-FCS artifact-control analyses, "
        "donor-level single-cell context, pathway context, and explicit claim-boundary auditing. These analyses "
        "strengthen the evidence chain but do not replace the requirement for prospective biological validation. "
        "The resulting bounded working interpretation is summarized in Figure 6, with the evidence hierarchy, "
        "claim-boundary matrix and reproducibility traceability summarized in Supplementary Figures S1 and S2."
    ),
    (
        "Artifact boundary for reviewers: throughout this manuscript, the residual CD3/CD14 co-event signal is "
        "treated as an artifact-aware event-derived summary, not as evidence for structural interaction biology "
        "beyond the measured event-derived signal. Matched raw-FCS negative-pair and all-pair marker-null "
        "sensitivity was feasible for the COMBAT sepsis subset and showed that pair specificity was not fully "
        "resolved, reinforcing frequency/state separation and QC-bound interpretation rather than a structural "
        "interaction claim."
    ): (
        "Artifact-aware interpretation. Throughout this manuscript, the representative residual co-event signal "
        "is treated as an artifact-aware event-derived summary within a broader marker-pair co-event background, "
        "not as evidence for structural interaction biology beyond the measured event-derived signal. Matched "
        "raw-FCS negative-pair and all-pair marker-null sensitivity was feasible for the COMBAT sepsis subset and "
        "showed that pair specificity was not fully resolved, reinforcing frequency/state separation and QC-bound "
        "interpretation rather than a structural interaction claim."
    ),
    "COMBAT Death28 was used only as an exploratory clinical anchor.": (
        "COMBAT 28-day death status (Death28) was used only as an exploratory clinical anchor."
    ),
    "Death28 was retained only as an exploratory clinical anchor": (
        "28-day death status (Death28) was retained only as an exploratory clinical anchor"
    ),
    "COMBAT Sepsis-only": "COMBAT sepsis-only",
    "40 Sepsis-only": "40 sepsis-only",
    "sepsis-only samples": "sepsis-only samples",
}

FIGURE_LEGENDS = {
    "Figure 1.": (
        "Figure 1. Study architecture and claim-boundary map. (A) Evidence hierarchy from public bulk RNA "
        "recurrence, continuous case-only immune spectrum, representative residual co-event analysis, paired "
        "COMBAT RNA-CyTOF bridging, donor-level single-cell context and exploratory 28-day death-status "
        "anchoring. (B) Locked interpretation: inflammatory/MHC-II decoupling is associated with a representative "
        "residual co-event signal within a broader marker-pair co-event background in paired sepsis samples. "
        "(C) Claim boundaries separating allowed cross-cohort, paired, activation-state and computational "
        "interpretations from excluded causal mechanism, structural interaction, validated endotype, clinical "
        "prediction and therapeutic claims."
    ),
    "Figure 4.": (
        "Figure 4. Raw-FCS marker-pair sensitivity and artifact-aware interpretation. (A) Prespecified marker-pair "
        "correlations across HLA-DR core, MHC-II/CD74 and RNA decoupling index. (B) All-pair raw-FCS marker-null "
        "distributions with CD3/CD14 positions. (C) Abundance-matched control-pair comparisons. (D) Lineage and "
        "QC residualized sensitivity. CD3/CD14 remains a representative residual co-event signal within a broader "
        "marker-pair co-event background; pair-level specificity remains unresolved, so the result is interpreted "
        "as an artifact-aware event-derived summary rather than a structural interaction."
    ),
    "Figure 5.": (
        "Figure 5. COMBAT paired RNA-CyTOF bridge. (A) Pairing workflow from RNA-seq logCPM and strict event-QC "
        "CyTOF summaries to 129 matched participant-timepoint rows and 40 sepsis-only paired rows from 34 "
        "participants. (B-D) Sepsis-only paired associations linking a representative residual co-event abundance "
        "metric, operationalized by the prespecified CD3/CD14 pair, to MHC-II/CD74, HLA-DR core and RNA "
        "decoupling index. (E) Bootstrap confidence intervals. (F) RNA-CyTOF pairing-permutation null "
        "distributions. (G) Participant-aware sensitivity. The figure supports a paired public-data bridge while "
        "preserving association-only, non-causal and non-predictive claim boundaries."
    ),
    "Figure 6.": (
        "Figure 6. Compartment-aware RNA-CyTOF immune-state decoupling map. (A) Whole-blood RNA layer summarizing "
        "inflammatory/S100 host-response increase, MHC-II/CD74/HLA-DR decrease and the decoupling index. (B) "
        "Paired COMBAT RNA-CyTOF bridge in sepsis-only samples. (C) CyTOF event-derived layer showing a "
        "representative residual co-event activation-state context with unresolved pair specificity. (D) Claim "
        "boundary: the diagram is an association-only, artifact-aware working interpretation requiring prospective "
        "validation, not a causal mechanism, clinical validation or structural interaction model."
    ),
}

PREFIX_REPLACEMENTS = {
    "Here we investigated whether public RNA and cytometry data could support": PARAGRAPH_REPLACEMENTS[
        (
            "Here we investigated whether public RNA and cytometry data could support a bounded, compartment-aware "
            "model of sepsis-associated immunoparalysis. The study architecture and claim boundaries are summarized "
            "in Figure 1. Our evidence hierarchy proceeds from cross-cohort bulk RNA recurrence to a continuous "
            "case-only immune-decoupling spectrum, raw-FCS pair-specificity hardening of the residual CD3/CD14 "
            "co-event signal, a paired COMBAT RNA-CyTOF bridge, donor-level single-cell contextual support, pathway "
            "context and an exploratory 28-day mortality clinical anchor. The central conclusion is constrained to "
            "the following: public RNA and cytometry data support a compartment-aware model of sepsis-associated "
            "immunoparalysis, in which systemic inflammatory/MHC-II decoupling is linked to a prespecified residual "
            "CD3/CD14 co-event signal in paired sepsis samples, while the event-derived signal itself remains an "
            "artifact-aware summary rather than a specific structural interaction or validated cell subset."
        )
    ],
    "This was a public-data systems immunology study designed for submission": PARAGRAPH_REPLACEMENTS[
        (
            "This was a public-data systems immunology study designed for submission to Frontiers in Immunology, "
            "Systems Immunology section, as an Original Research article. No MIMIC-IV, eICU, protected health "
            "information, new human samples, wet-laboratory experiments, MR/SMR causal analysis, clinical "
            "prediction modeling, drug repositioning, or treatment-response modeling was performed."
        )
    ],
    "Cytometry summaries were interpreted with explicit artifact boundaries.": PARAGRAPH_REPLACEMENTS[
        (
            "Cytometry summaries were interpreted with explicit artifact boundaries. Frequency and summary-level "
            "marker-state context were treated as distinct readouts. Matching public raw whole-blood FCS files were "
            "available for all 40 COMBAT sepsis-only paired rows and enabled event-level negative-pair and "
            "pair-specificity hardening. In the required marker-pair comparison, CD3/CD14 showed the expected "
            "association direction but was not uniquely strongest: for the RNA decoupling index, CD14/CD16 and "
            "CD4/CD14 showed larger rho values than CD3/CD14, whereas CD19/CD14, CD8/CD14 and CD56/CD14 were weaker "
            "or opposite. In the all-pair raw-FCS marker-null analysis of 378 available marker pairs, CD3/CD14 "
            "ranked in the 61.6% to 65.1% percentile range by absolute rho across the three RNA metrics, and 4 to 5 "
            "of 12 abundance-matched control pairs showed absolute correlations at least as large as CD3/CD14. "
            "Threshold perturbation and technical-covariate residualization preserved the expected direction for "
            "CD3/CD14, whereas event-count downsampling weakened this pattern. These results reduce the earlier "
            "feasibility gap but leave CD3/CD14 pair-level specificity unresolved; CD3/CD14 is therefore retained "
            "as a prespecified representative residual co-event signal, not as evidence for a specific structural "
            "interaction or new cell subset (Figure 4; Supplementary Tables S8-S10 and S19)."
        )
    ],
    "Within sepsis residual CD3/CD14 events, summary-level marker-state context included": PARAGRAPH_REPLACEMENTS[
        (
            "Within sepsis residual CD3/CD14 events, summary-level marker-state context included HLA-DR positivity "
            "(median 0.100, IQR 0.029-0.223), CD33 positivity (median 0.073, IQR 0.014-0.142), CD11c positivity "
            "(median 0.067, IQR 0.009-0.202) and CD38 positivity (median 0.317, IQR 0.186-0.494). These summaries "
            "make the activation-state remodeling more concrete, but they do not define a new cell subset or prove "
            "structural contact."
        )
    ],
    "COMBAT paired analysis provided the most direct paired public-data bridge": PARAGRAPH_REPLACEMENTS[
        (
            "COMBAT paired analysis provided the most direct paired public-data bridge in this study. The paired "
            "dataset contained 129 matched biological RNA-CyTOF participant-timepoint rows, including 40 sepsis-only "
            "rows from 34 participants. Six participants contributed repeated sepsis timepoints, with a maximum of "
            "two rows per participant. In sepsis-only samples, abundance-normalized residual CD3/CD14 co-event signal "
            "was inversely associated with HLA-DR core RNA score and MHC-II/CD74 RNA score, and positively associated "
            "with the RNA decoupling index. Row-level correlations, participant-level averaged correlations, "
            "participant-cluster bootstrap, and participant-cluster rank-regression sensitivity preserved the same "
            "association directions, although uncertainty increased after collapsing or resampling repeated "
            "observations (Figure 5; Table 5; Supplementary Table S16)."
        )
    ],
    "This study supports a bounded systems immunology reconstruction": PARAGRAPH_REPLACEMENTS[
        (
            "This study supports a bounded systems immunology reconstruction of sepsis-associated immunoparalysis. "
            "The fundamental contribution is not the demonstration that MHC-II or HLA-DR suppression exists in "
            "sepsis-this is already well established-but rather that systemic inflammatory/MHC-II decoupling can be "
            "linked to a cytometry-defined residual CD3/CD14 co-event signal in matched sepsis samples, while the "
            "residual event-derived summary itself exhibits activation-state remodeling."
        )
    ],
    "The cytometry evidence also necessitates conservative reformulation.": PARAGRAPH_REPLACEMENTS[
        (
            "The cytometry evidence also necessitates conservative reformulation. Raw-FCS pair-specificity hardening "
            "did not establish CD3/CD14 as a uniquely specific pair; rather, CD3/CD14 remains a prespecified "
            "representative residual co-event signal within a broader co-event background. The safer interpretation "
            "is one of analytical-layer divergence: whole-blood RNA can capture population-level antigen-presentation "
            "suppression while residual event-derived summaries show summary-level marker-state context and co-event "
            "frequency behavior that are not reducible to a single structural interaction claim."
        )
    ],
    "The stronger CD14/CD16 and CD4/CD14 associations provide": PARAGRAPH_REPLACEMENTS[
        (
            "The stronger CD14/CD16 and CD4/CD14 associations provide a biologically plausible caution against "
            "over-centering a single pair: CD14/CD16 may capture broader myeloid and monocyte-lineage co-event "
            "background, whereas CD4/CD14 may be closer to helper T-cell/monocyte antigen-presentation context. "
            "Accordingly, the study treats CD3/CD14 as a prespecified representative co-event signal rather than as "
            "a uniquely specific pair."
        )
    ],
    "The validation strategy reflects the inherent constraints": PARAGRAPH_REPLACEMENTS[
        (
            "The validation strategy reflects the inherent constraints of public-data computational research. "
            "Frontiers-relevant validation is addressed through cross-cohort recurrence, continuous-spectrum analysis, "
            "broad-cell composition sensitivity, orthogonal paired RNA-CyTOF bridging, participant-level robustness, "
            "raw-FCS artifact-control hardening, single-cell donor-level contextual analysis, pathway enrichment "
            "context and explicit claim-boundary auditing. These analyses strengthen the evidence chain but do not "
            "replace the requirement for prospective biological validation. The resulting bounded working "
            "interpretation is summarized in Figure 6, with the evidence hierarchy, claim-boundary matrix and "
            "reproducibility traceability summarized in Supplementary Figures S1 and S2."
        )
    ],
    "Artifact boundary for reviewers:": PARAGRAPH_REPLACEMENTS[
        (
            "Artifact boundary for reviewers: throughout this manuscript, the residual CD3/CD14 co-event signal is "
            "treated as an artifact-aware event-derived summary, not as evidence for structural interaction biology "
            "beyond the measured event-derived signal. Matched raw-FCS negative-pair and all-pair marker-null "
            "sensitivity was feasible for the COMBAT sepsis subset and showed that pair specificity was not fully "
            "resolved, reinforcing frequency/state separation and QC-bound interpretation rather than a structural "
            "interaction claim."
        )
    ],
}


def run(cmd: list[str], *, check: bool = True, env: dict[str, str] | None = None) -> subprocess.CompletedProcess[str]:
    proc = subprocess.run(
        cmd,
        cwd=str(ROOT),
        env=env,
        check=False,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    if check and proc.returncode != 0:
        raise RuntimeError(
            "Command failed:\n"
            + " ".join(cmd)
            + f"\nexit={proc.returncode}\nSTDOUT:\n{proc.stdout}\nSTDERR:\n{proc.stderr}"
        )
    return proc


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def rel(path: Path) -> str:
    try:
        return path.relative_to(ROOT).as_posix()
    except ValueError:
        return str(path)


def clear_set(paragraph, text: str) -> None:
    paragraph.clear()
    run_obj = paragraph.add_run(text)
    run_obj.font.name = "Arial"
    run_obj._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run_obj._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")


def iter_paragraphs_including_tables(doc: Document):
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def citation_payload_to_parentheses(payload: str) -> str:
    payload = payload.replace(" ", "")
    parts = []
    for part in payload.split(","):
        parts.append(part.replace("-", "–"))
    return "(" + ",".join(parts) + ")"


def convert_citations(text: str) -> str:
    return re.sub(r"\[(\d+(?:-\d+)?(?:,\d+(?:-\d+)?)*)\]", lambda m: citation_payload_to_parentheses(m.group(1)), text)


def convert_reference_style(text: str) -> str:
    if text.startswith("9. Snow TAC, Villa A, Cesar A, Ryckaert F, Saleem N, Smyth D, et al."):
        return (
            "9. Snow TAC, Villa A, Cesar A, Ryckaert F, Saleem N, Smyth D, et al. "
            "Challenges of monocyte HLA-DR targeted immunomodulation in sepsis—a prospective observational "
            "cohort study. Front Immunol (2026) 16:1709289. doi: 10.3389/fimmu.2025.1709289."
        )
    match = re.match(r"^(\d+)\.\s+(.+)\.\s+([^\.]+)\.\s+(\d{4});([^\.]+)\.\s+doi:\s+(.+)\.$", text)
    if not match:
        return text
    number, before, journal, year, volpages, doi = match.groups()
    return f"{number}. {before}. {journal} ({year}) {volpages}. doi: {doi}."


def patch_main_docx() -> dict[str, object]:
    if not SOURCE_MAIN.exists():
        raise FileNotFoundError(SOURCE_MAIN)
    doc = Document(SOURCE_MAIN)
    hits: dict[str, int] = {}
    in_references = False

    for para_idx, para in enumerate(doc.paragraphs):
        text = " ".join(para.text.split())
        if not text:
            continue
        original = text
        if text.startswith("Compartment-aware reconstruction"):
            text = TITLE
        elif text == "Contribution to the Field":
            pass
        elif para_idx == 8:
            text = CONTRIBUTION
        elif text.startswith("Background: Whole-blood RNA sequencing"):
            text = ABSTRACT
        elif text == "sepsis, systems immunology, immunoparalysis, MHC-II, HLA-DR, CyTOF, CD3/CD14 co-event, RNA-CyTOF":
            text = "sepsis, systems immunology, immunoparalysis, antigen presentation, cytometry co-event signal, RNA-CyTOF"
        elif text == "References":
            in_references = True
        elif text == "Figure Legends":
            in_references = False
        elif in_references and re.match(r"^\d+\.\s+", text):
            text = convert_reference_style(text)
        else:
            for prefix, new_text in PREFIX_REPLACEMENTS.items():
                if text.startswith(prefix):
                    text = new_text
                    hits[prefix] = hits.get(prefix, 0) + 1
                    break
            for old, new in PARAGRAPH_REPLACEMENTS.items():
                if old in text:
                    text = text.replace(old, new)
                    hits[old] = hits.get(old, 0) + 1
            for prefix, new_text in FIGURE_LEGENDS.items():
                if text.startswith(prefix):
                    text = new_text
                    hits[prefix] = hits.get(prefix, 0) + 1
                    break
            if text.startswith("In the sepsis-only paired subset, Death28 involved"):
                text = text.replace(
                    "In the sepsis-only paired subset, Death28 involved",
                    "In the sepsis-only paired subset, Death28 involved",
                )
            if not in_references:
                text = convert_citations(text)
        if text != original:
            clear_set(para, text)

    for para in iter_paragraphs_including_tables(doc):
        text = para.text
        new = text.replace("COMBAT Sepsis-only", "COMBAT sepsis-only")
        new = new.replace("40 Sepsis-only", "40 sepsis-only")
        new = new.replace("Supplementary Table S17. Death28 exploratory clinical anchor", "Supplementary Table S17. 28-day death status (Death28) exploratory clinical anchor")
        if new != text:
            clear_set(para, new)

    doc.save(MAIN_DOCX)
    UPLOAD.mkdir(parents=True, exist_ok=True)
    shutil.copy2(MAIN_DOCX, UPLOAD / MAIN_DOCX.name)
    return {
        "source": rel(SOURCE_MAIN),
        "docx": rel(MAIN_DOCX),
        "upload_docx": rel(UPLOAD / MAIN_DOCX.name),
        "replacement_hit_count": sum(hits.values()),
        "replacement_hit_categories": len(hits),
        "contribution_words": len(CONTRIBUTION.split()),
        "abstract_words": len(ABSTRACT.split()),
    }


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def find_source_file(filename: str) -> Path | None:
    candidates = [p for p in ROOT.rglob(filename) if p.is_file()]
    candidates = [p for p in candidates if "07_FI_SUBMISSION_REVISED_20260622" not in p.as_posix()]
    if not candidates:
        return None
    candidates.sort(key=lambda p: (0 if "03_TRACEABLE_RESULTS" in p.as_posix() else 1, len(p.as_posix())))
    return candidates[0]


def csv_shape(path: Path) -> tuple[int, int, list[str]]:
    try:
        with path.open(newline="", encoding="utf-8-sig") as handle:
            reader = csv.reader(handle)
            header = next(reader, [])
            rows = sum(1 for _ in reader)
        return rows, len(header), header
    except Exception:
        return 0, 0, []


def slug_name(text: str) -> str:
    slug = re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_")
    return slug[:42] or "table"


def clean_supplement_title(title: str) -> str:
    cleaned = title.replace("Sepsis-only", "sepsis-only")
    cleaned = cleaned.replace("Death28 clinical anchor", "28-day death status (Death28) clinical anchor")
    cleaned = cleaned.replace("cytometry marker-state remodeling results", "cytometry marker-state heterogeneity results")
    return cleaned


def sanitize_output_text(text: str) -> str:
    replacements = {
        "COMBAT Sepsis-only": "COMBAT sepsis-only",
        "Sepsis-only": "sepsis-only",
        "11 Death28 events": "11 28-day death-status events",
        "associated with Death28": "associated with 28-day death status (Death28)",
        "Death28 clinical anchor": "28-day death status (Death28) clinical anchor",
        "Residual CD3/CD14 co-events show activation-state remodeling": "Residual co-event summaries show activation-state context",
        "activation-state remodeling": "activation-state context",
        "marker-state remodeling": "marker-state heterogeneity",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def sanitize_text_file(path: Path) -> None:
    text = path.read_text(encoding="utf-8", errors="ignore")
    cleaned = sanitize_output_text(text)
    if cleaned != text:
        path.write_text(cleaned, encoding="utf-8")


def build_supplementary_csv_package() -> tuple[list[dict[str, object]], Path]:
    if not SUPP_INDEX.exists():
        raise FileNotFoundError(SUPP_INDEX)
    if TABLE_DIR.exists():
        shutil.rmtree(TABLE_DIR)
    TABLE_DIR.mkdir(parents=True, exist_ok=True)

    index_rows = read_csv_rows(SUPP_INDEX)
    package_rows: list[dict[str, object]] = []
    for row in index_rows:
        s_label = row["supplementary_table"].strip()
        title = clean_supplement_title(row["title"].strip())
        sources = [item.strip() for item in row["source_or_status"].split(";") if item.strip()]
        for source_idx, source_name in enumerate(sources, start=1):
            source_path = find_source_file(source_name)
            if source_path is None:
                package_rows.append(
                    {
                        "supplementary_table": s_label,
                        "title": title,
                        "source_name": source_name,
                        "copied_file": "",
                        "source_found": False,
                        "rows": "",
                        "columns": "",
                        "sha256": "",
                    }
                )
                continue
            s_number = int(re.sub(r"\D", "", s_label) or "0")
            s_code = f"S{s_number:02d}"
            suffix = f"_part{source_idx}" if len(sources) > 1 else ""
            dest_name = f"{s_code}{suffix}_{slug_name(title)}.csv"
            dest = TABLE_DIR / dest_name
            shutil.copy2(source_path, dest)
            sanitize_text_file(dest)
            rows, cols, header = csv_shape(dest)
            package_rows.append(
                {
                    "supplementary_table": s_label,
                    "title": title,
                    "source_name": source_name,
                    "copied_file": dest.name,
                    "source_found": True,
                    "rows": rows,
                    "columns": cols,
                    "column_preview": "; ".join(header[:8]),
                    "sha256": sha256(dest),
                }
            )

    manifest = TABLE_DIR / "Supplementary_Tables_S1_S22_CSV_manifest.csv"
    write_csv(manifest, package_rows)
    zip_path = UPLOAD / "Supplementary_Tables_S1_S22_CSV.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(TABLE_DIR.glob("*.csv")):
            zf.write(path, f"Supplementary_Tables_S1_S22_CSV/{path.name}")
    return package_rows, zip_path


def set_normal_font(doc: Document) -> None:
    styles = doc.styles
    for name, size in [("Normal", 9), ("Heading 1", 14), ("Heading 2", 11), ("Heading 3", 10)]:
        if name in styles:
            styles[name].font.name = "Arial"
            styles[name].font.size = Pt(size)


def set_landscape(doc: Document) -> None:
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 7.2) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run_obj = p.add_run(str(text))
    run_obj.font.name = "Arial"
    run_obj.font.size = Pt(size)
    run_obj.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float], font_size: float = 7.0) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=font_size)
        shade_cell(table.rows[0].cells[i], "EAF2F8")
        table.rows[0].cells[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
            cells[i].width = Inches(widths[i])
    doc.add_paragraph()


def add_preview_table(doc: Document, csv_path: Path, max_rows: int = 4, max_cols: int = 5) -> None:
    rows = []
    with csv_path.open(newline="", encoding="utf-8-sig") as handle:
        reader = csv.reader(handle)
        header = next(reader, [])
        for idx, row in enumerate(reader):
            if idx >= max_rows:
                break
            rows.append(row)
    if not header:
        return
    header = header[:max_cols]
    clipped = [[cell[:90] for cell in row[:max_cols]] for row in rows]
    widths = [max(1.2, min(2.2, 9.6 / len(header)))] * len(header)
    add_table(doc, header, clipped, widths, font_size=6.1)


def build_supplementary_index_docx(package_rows: list[dict[str, object]], zip_path: Path) -> dict[str, object]:
    doc = Document()
    set_landscape(doc)
    set_normal_font(doc)
    doc.add_heading("Supplementary Materials Index", level=1)
    p = doc.add_paragraph()
    p.add_run("Scope. ").bold = True
    p.add_run(
        "This file intentionally contains only a Supplementary Table Index and short, narrow previews. "
        "Complete wide tables are supplied as editable CSV files in the companion zip archive."
    )
    doc.add_paragraph(f"Editable table package: {zip_path.name}")

    index_rows: list[list[object]] = []
    by_s: dict[str, list[dict[str, object]]] = {}
    for row in package_rows:
        by_s.setdefault(str(row["supplementary_table"]), []).append(row)
    for s_label in sorted(by_s, key=lambda s: int(re.sub(r"\D", "", s) or "0")):
        rows = by_s[s_label]
        files = "; ".join(str(r["copied_file"]) for r in rows if r.get("copied_file"))
        row_count = "; ".join(str(r["rows"]) for r in rows if r.get("rows") != "")
        col_count = "; ".join(str(r["columns"]) for r in rows if r.get("columns") != "")
        index_rows.append([s_label, rows[0]["title"], files, row_count, col_count])
    add_table(
        doc,
        ["Table", "Title", "Editable CSV file(s)", "Rows", "Cols"],
        index_rows,
        [0.65, 2.35, 5.5, 0.7, 0.55],
        font_size=6.2,
    )

    doc.add_heading("Short Preview Tables", level=1)
    preview_labels = {"S17", "S19", "S21"}
    for s_label in sorted(preview_labels, key=lambda s: int(s[1:])):
        rows = by_s.get(s_label, [])
        if not rows:
            continue
        doc.add_heading(f"{s_label}. {rows[0]['title']}", level=2)
        doc.add_paragraph(
            "Preview only. The full editable table is in the CSV zip archive; columns and rows are intentionally capped here for PDF readability."
        )
        for item in rows[:2]:
            copied = item.get("copied_file")
            if not copied:
                continue
            doc.add_paragraph(str(copied))
            add_preview_table(doc, TABLE_DIR / str(copied))

    doc.add_heading("Submission Handling Note", level=1)
    doc.add_paragraph(
        "Large tables should be uploaded as editable supplementary material rather than forced into the manuscript PDF. "
        "The main manuscript contains centralized figure legends and no embedded figure images; Figure 1-6 and Supplementary "
        "Figures S1-S2 are provided separately as upload-ready TIFF files."
    )
    doc.save(SUPP_DOCX)
    shutil.copy2(SUPP_DOCX, UPLOAD / SUPP_DOCX.name)
    return {"docx": rel(SUPP_DOCX), "upload_docx": rel(UPLOAD / SUPP_DOCX.name)}


def render_docx(docx_path: Path, variant: str) -> dict[str, object]:
    out_dir = QA / f"render_{variant}"
    if out_dir.exists():
        shutil.rmtree(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    env = os.environ.copy()
    if SOFFICE_APP.exists():
        env["PATH"] = f"{SOFFICE_APP}:{env.get('PATH', '')}"
    proc = run(
        [sys.executable, str(DOCX_RENDERER), str(docx_path), "--output_dir", str(out_dir), "--emit_pdf"],
        check=False,
        env=env,
    )
    pdfs = sorted(out_dir.glob("*.pdf"))
    pngs = sorted(out_dir.glob("page-*.png"))
    result = {
        "variant": variant,
        "docx": rel(docx_path),
        "success": proc.returncode == 0 and bool(pdfs) and bool(pngs),
        "returncode": proc.returncode,
        "rendered_png_count": len(pngs),
        "stdout_tail": proc.stdout[-1200:],
        "stderr_tail": proc.stderr[-1200:],
        "png_dir": rel(out_dir),
    }
    if result["success"]:
        pdf_out = UPLOAD / f"{docx_path.stem}.pdf"
        shutil.copy2(pdfs[0], pdf_out)
        result["pdf"] = rel(pdf_out)
        result["page_count"] = len(pngs)
    return result


def make_contact_sheet(render_dir: Path, out_path: Path) -> None:
    pages = sorted(render_dir.glob("page-*.png"), key=lambda p: int(p.stem.split("-")[1]))
    thumbs = []
    for page in pages:
        img = Image.open(page).convert("RGB")
        img.thumbnail((260, 200))
        tile = Image.new("RGB", (290, 238), "white")
        tile.paste(img, ((290 - img.width) // 2, 8))
        thumbs.append((tile, page.stem))
    cols = 3
    rows = max(1, (len(thumbs) + cols - 1) // cols)
    sheet = Image.new("RGB", (cols * 290, rows * 238), "white")
    draw = ImageDraw.Draw(sheet)
    for idx, (tile, label) in enumerate(thumbs):
        x = (idx % cols) * 290
        y = (idx // cols) * 238
        sheet.paste(tile, (x, y))
        draw.text((x + 8, y + 216), label, fill=(0, 0, 0))
    sheet.save(out_path, dpi=(150, 150))


def count_docx_drawings(docx_path: Path) -> dict[str, object]:
    with zipfile.ZipFile(docx_path) as zf:
        names = zf.namelist()
        xml = "\n".join(
            zf.read(name).decode("utf-8", errors="ignore")
            for name in names
            if name.startswith("word/") and name.endswith(".xml")
        )
        media = [name for name in names if name.startswith("word/media/")]
    return {
        "drawing_tags": len(re.findall(r"<w:drawing[\s>]", xml)),
        "media_count": len(media),
        "tracked_change_tags": sum(
            len(re.findall(pattern, xml))
            for pattern in (r"<w:ins[\s>]", r"<w:del[\s>]", r"<w:moveFrom[\s>]", r"<w:moveTo[\s>]")
        ),
        "comment_parts": [name for name in names if name.startswith("word/comments")],
    }


def doc_text(docx_path: Path) -> str:
    doc = Document(docx_path)
    parts = []
    for para in iter_paragraphs_including_tables(doc):
        parts.append(para.text)
    return "\n".join(parts)


def manuscript_qa(docx_path: Path) -> dict[str, object]:
    text = doc_text(docx_path)
    forbidden = {
        "journal_submission_trace": "designed for submission to Frontiers",
        "journal_named_validation": "Frontiers-relevant validation",
        "reviewer_facing_artifact_label": "Artifact boundary for reviewers",
        "missing_sepsis_comma": "sepsis-this",
        "missing_pair_space": "specific pair.Clinically",
        "capitalized_sepsis_only": "COMBAT Sepsis-only",
        "local_summary_ethics_phrase": "local pre-analyzed derived summaries",
        "strong_activation_remodeling": "activation-state remodeling",
        "strong_coevent_remodeling": "CD3/CD14 co-event remodeling",
        "linked_immunoparalysis_phrase": "CD3/CD14-linked immunoparalysis",
        "specific_signal_phrase": "CD3/CD14-specific signal",
    }
    citation_square_hits = re.findall(r"\[\d+(?:-\d+)?(?:,\d+(?:-\d+)?)*\]", text)
    return {
        "docx": rel(docx_path),
        "sha256": sha256(docx_path),
        "forbidden_hits": {label: text.count(term) for label, term in forbidden.items()},
        "square_numeric_citation_hits": citation_square_hits[:20],
        "square_numeric_citation_count": len(citation_square_hits),
        "contribution_words": len(CONTRIBUTION.split()),
        "abstract_words": len(ABSTRACT.split()),
        **count_docx_drawings(docx_path),
    }


def write_csv(path: Path, rows: Iterable[dict[str, object]]) -> None:
    rows = list(rows)
    keys: list[str] = []
    for row in rows:
        for key in row:
            if key not in keys:
                keys.append(key)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=keys)
        writer.writeheader()
        writer.writerows(rows)


def copy_upload_figures() -> list[dict[str, object]]:
    rows = []
    for name in FIGURE_FILES + PREVIEW_FILES:
        src = FIGURE_FILE_OVERRIDES.get(name, OLD_UPLOAD / name)
        if not src.exists():
            rows.append({"file": name, "copied": False, "source": rel(src), "sha256": ""})
            continue
        dst = UPLOAD / name
        shutil.copy2(src, dst)
        rows.append({"file": name, "copied": True, "source": rel(src), "dest": rel(dst), "sha256": sha256(dst)})
    return rows


def write_manifest() -> None:
    rows = []
    for path in sorted(p for p in OUT.rglob("*") if p.is_file()):
        rows.append({"relative_path": path.relative_to(OUT).as_posix(), "size_bytes": path.stat().st_size, "sha256": sha256(path)})
    write_csv(QA / "FI_submission_revised_manifest_sha256.csv", rows)


def write_readme(zip_path: Path, render_results: list[dict[str, object]]) -> None:
    lines = [
        "# FI submission revised package",
        "",
        f"Generated: {datetime.now().isoformat(timespec='seconds')}",
        "",
        "## Primary Upload Files",
        f"- `01_upload_ready/{MAIN_DOCX.name}`: revised main manuscript, text and centralized figure legends only; no embedded figure images.",
        f"- `01_upload_ready/{MAIN_DOCX.stem}.pdf`: rendered manuscript PDF from the revised main manuscript.",
        f"- `01_upload_ready/{SUPP_DOCX.name}`: supplementary-materials index with short previews only.",
        f"- `01_upload_ready/{SUPP_DOCX.stem}.pdf`: rendered supplementary-materials index PDF.",
        "- `01_upload_ready/Figure_01.tiff` through `Figure_06.tiff`: separate upload-ready main figure files.",
        "- `01_upload_ready/Supplementary_Figure_S1.tiff` and `Supplementary_Figure_S2.tiff`: separate upload-ready supplementary figures.",
        f"- `01_upload_ready/{zip_path.name}`: editable CSV supplement package for Supplementary Tables S1-S22.",
        "",
        "## Revision Notes",
        "- Removed target-journal/submission-trace wording from the manuscript body.",
        "- Reframed CD3/CD14 as a representative residual co-event signal within a broader marker-pair co-event background.",
        "- Replaced reviewer-facing wording with artifact-aware interpretation.",
        "- Converted numeric in-text citations from square brackets to parentheses and changed references to a Frontiers-style year-in-parentheses format.",
        "- Removed embedded figure images from the primary manuscript to avoid caption/figure page splits; figures are supplied separately as TIFF.",
        "- Kept wide supplementary tables out of the PDF and supplied them as editable CSV files.",
        "",
        "## Render QA",
    ]
    for item in render_results:
        lines.append(f"- {item['variant']}: success={item.get('success')} pages={item.get('page_count', item.get('rendered_png_count'))}")
    (OUT / "README.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def zip_revised_package() -> dict[str, object]:
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(p for p in OUT.rglob("*") if p.is_file()):
            zf.write(path, path.relative_to(ROOT).as_posix())
    proc = run(["unzip", "-t", str(ZIP_PATH)], check=False)
    (QA / "zip_test.txt").write_text(proc.stdout + proc.stderr, encoding="utf-8")
    return {"zip": rel(ZIP_PATH), "success": proc.returncode == 0, "returncode": proc.returncode, "sha256": sha256(ZIP_PATH), "size_bytes": ZIP_PATH.stat().st_size}


def main() -> None:
    if OUT.exists():
        shutil.rmtree(OUT)
    UPLOAD.mkdir(parents=True, exist_ok=True)
    QA.mkdir(parents=True, exist_ok=True)
    for path in [SOURCE_MAIN, DOCX_RENDERER, SUPP_INDEX]:
        if not path.exists():
            raise FileNotFoundError(path)

    main_result = patch_main_docx()
    figure_copy = copy_upload_figures()
    package_rows, table_zip = build_supplementary_csv_package()
    supp_result = build_supplementary_index_docx(package_rows, table_zip)

    render_results = [
        render_docx(MAIN_DOCX, "v21_main_text_only_revised"),
        render_docx(SUPP_DOCX, "v21_supplementary_index_only"),
    ]
    for item in render_results:
        if item["success"]:
            make_contact_sheet(ROOT / item["png_dir"], QA / f"{item['variant']}_contact_sheet.png")

    package_manifest = QA / "Supplementary_Tables_S1_S22_CSV_manifest.csv"
    write_csv(package_manifest, package_rows)
    write_manifest()
    zip_result = zip_revised_package()
    if not zip_result["success"]:
        raise RuntimeError("Revised package zip failed unzip -t")
    write_readme(table_zip, render_results)
    write_manifest()

    report = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "main_revision": main_result,
        "supplementary_index": supp_result,
        "figure_copy": figure_copy,
        "supplementary_table_zip": {"path": rel(table_zip), "sha256": sha256(table_zip), "size_bytes": table_zip.stat().st_size},
        "supplementary_table_manifest": rel(package_manifest),
        "manuscript_QA": manuscript_qa(MAIN_DOCX),
        "supplementary_QA": manuscript_qa(SUPP_DOCX),
        "render_QA": render_results,
        "zip_result": zip_result,
        "frontiers_ref9_source_checked": "https://www.frontiersin.org/journals/immunology/articles/10.3389/fimmu.2025.1709289/full",
    }
    (QA / "revision_QA_report.json").write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps(
        {
            "main_docx": str(MAIN_DOCX),
            "supplementary_index_docx": str(SUPP_DOCX),
            "output_dir": str(OUT),
            "table_zip": str(table_zip),
            "package_zip": str(ZIP_PATH),
            "render_success": [item["success"] for item in render_results],
            "zip_success": zip_result["success"],
        },
        indent=2,
    ))


if __name__ == "__main__":
    main()
